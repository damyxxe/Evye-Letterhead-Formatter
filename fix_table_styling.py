"""
Modernise fee table styling:
1. Remove grid lines — use subtle grey separators only
2. Add generous cell padding (3.5mm vertical)
3. Light grey header background (#F5F5F5)
4. Right-align Fee column, accounting-style numbers
5. Clean total row with single top rule
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_ROOT = Path(__file__).resolve().parent
TEMPLATE_PATH = PROJECT_ROOT / "templates" / "evye-master.docx"


def fix_fee_table(doc):
    """Find and restyle the fee table."""
    # Find the fee table (first table in the document body)
    if not doc.tables:
        print("  No tables found!")
        return

    table = doc.tables[0]
    print(f"  Found table: {len(table.rows)} rows x {len(table.columns)} cols")

    # ── 1. Remove ALL table-level borders ──
    _set_table_borders_none(table)

    # ── 2. Style each row ──
    for i, row in enumerate(table.rows):
        if i == 0:
            # Header row: light grey background, thin bottom border
            _style_header_row(row)
        elif i == len(table.rows) - 1:
            # Total row: thin top rule, no bottom
            _style_total_row(row, table)
        else:
            # Data rows: hairline bottom border only
            _style_data_row(row)

        # ── 3. Set cell padding on all rows ──
        for cell in row.cells:
            _set_cell_margins(cell, top=Mm(1.5), bottom=Mm(1.5), left=Mm(1.5), right=Mm(1.5))

    # ── 4. Right-align Fee column (last column) and fix number format ──
    for i, row in enumerate(table.rows):
        fee_cell = row.cells[-1]  # Last column
        for p in fee_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Fix number formatting: "$ 4,500" → "$4,500"
            for run in p.runs:
                if run.text.strip().startswith('$'):
                    run.text = run.text.replace('$ ', '$')

    print("  Applied modern styling")


def _set_table_borders_none(table):
    """Remove all borders from the table element."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    # Remove existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _style_header_row(row):
    """Header: light grey background (#F5F5F5), thin bottom border."""
    for cell in row.cells:
        # Grey background
        _set_cell_shading(cell, "F5F5F5")
        # Thin bottom border only
        _set_cell_borders(cell,
                          top="none", left="none", right="none",
                          bottom=("single", "4", "999999"))


def _style_data_row(row):
    """Data rows: hairline bottom border in very light grey."""
    for cell in row.cells:
        _set_cell_shading(cell, "FFFFFF")  # White background
        _set_cell_borders(cell,
                          top="none", left="none", right="none",
                          bottom=("single", "2", "E5E5E5"))


def _style_total_row(row, table):
    """Total row: thin top rule, bold, no bottom border."""
    for cell in row.cells:
        _set_cell_shading(cell, "FFFFFF")
        _set_cell_borders(cell,
                          top=("single", "4", "000000"),
                          left="none", right="none", bottom="none")


def _set_cell_shading(cell, color):
    """Set cell background colour."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)

    # Remove existing shading
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)

    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
    )
    tcPr.append(shd)


def _set_cell_borders(cell, top="none", bottom="none", left="none", right="none"):
    """Set individual cell borders. Each param is either 'none' or (style, size, color)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)

    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)

    def border_xml(name, val):
        if val == "none":
            return f'<w:{name} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        style, sz, color = val
        return f'<w:{name} w:val="{style}" w:sz="{sz}" w:space="0" w:color="{color}"/>'

    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'{border_xml("top", top)}'
        f'{border_xml("bottom", bottom)}'
        f'{border_xml("left", left)}'
        f'{border_xml("right", right)}'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """Set cell margins (padding)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)

    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)

    margins = f'<w:tcMar {nsdecls("w")}>'
    if top is not None:
        margins += f'<w:top w:w="{int(top)}" w:type="dxa"/>'
    if bottom is not None:
        margins += f'<w:bottom w:w="{int(bottom)}" w:type="dxa"/>'
    if left is not None:
        margins += f'<w:left w:w="{int(left)}" w:type="dxa"/>'
    if right is not None:
        margins += f'<w:right w:w="{int(right)}" w:type="dxa"/>'
    margins += '</w:tcMar>'

    tcPr.append(parse_xml(margins))


def fix_table_styles(doc):
    """Update the table cell styles for right-aligned fees."""
    print("Updating Evye Table Total style to right-align...")
    try:
        tt = doc.styles["Evye Table Total"]
        tt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    except KeyError:
        pass


def main():
    print(f"Opening: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))

    print("Restyling fee table...")
    fix_fee_table(doc)
    fix_table_styles(doc)

    doc.save(str(TEMPLATE_PATH))
    print(f"\nSaved: {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
