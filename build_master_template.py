"""
Build the Evye master .docx template with all brand styles.

This creates evye-master.docx — the template that the generator clones
for every new document. Edit it in Word to fine-tune, then save.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

PROJECT_ROOT = Path(__file__).resolve().parent
ASSETS_DIR = PROJECT_ROOT / "assets"
FONTS_DIR = PROJECT_ROOT / "fonts"
TEMPLATES_DIR = PROJECT_ROOT / "templates"

# ── Brand constants ──────────────────────────────────────────

TELEGRAF = "Telegraf"
HELVETICA = "Helvetica Neue"
FALLBACK_BODY = "Arial"

EVYE_BLACK = RGBColor(0x00, 0x00, 0x00)
EVYE_GREY = RGBColor(0x99, 0x99, 0x99)
EVYE_DARK_GREY = RGBColor(0x33, 0x33, 0x33)
EVYE_LIGHT_GREY = RGBColor(0xCC, 0xCC, 0xCC)
EVYE_TABLE_BORDER = RGBColor(0xCC, 0xCC, 0xCC)

# A4 page with brand margins (matching CSS: 15mm top, 20mm left, 15mm right, 22mm bottom)
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_TOP = Mm(15)
MARGIN_BOTTOM = Mm(22)
MARGIN_LEFT = Mm(20)
MARGIN_RIGHT = Mm(15)  # Original right margin — content width controlled via paragraph indent


def create_master_template():
    doc = Document()

    # ── Page Setup ───────────────────────────────────────────
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # ── Clean up default styles ──────────────────────────────
    _setup_default_font(doc)

    # ── Create all brand styles ──────────────────────────────
    _create_heading_styles(doc)
    _create_body_styles(doc)
    _create_clause_styles(doc)
    _create_list_styles(doc)
    _create_table_styles(doc)
    _create_special_styles(doc)

    # ── Header with logo ─────────────────────────────────────
    _setup_header(doc, section)

    # ── Footer ───────────────────────────────────────────────
    _setup_footer(doc, section)

    # ── Sample content to demonstrate styles ─────────────────
    _add_sample_content(doc)

    # ── Save ─────────────────────────────────────────────────
    output_path = TEMPLATES_DIR / "evye-master.docx"
    doc.save(str(output_path))
    print(f"Master template created: {output_path}")
    return str(output_path)


def _setup_default_font(doc):
    """Set the default document font to Helvetica Neue."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = HELVETICA
    font.size = Pt(9.5)
    font.color.rgb = EVYE_BLACK
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35
    pf.right_indent = Mm(42)  # Keep body content clear of the header sidebar area

    # Set East Asian / complex script fallback
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{HELVETICA}" w:hAnsi="{HELVETICA}" w:cs="{FALLBACK_BODY}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), HELVETICA)
        rFonts.set(qn("w:hAnsi"), HELVETICA)
        rFonts.set(qn("w:cs"), FALLBACK_BODY)


def _create_heading_styles(doc):
    """Create heading hierarchy: Telegraf for H1/H2, Helvetica Neue for H3/H4."""

    # ── Document Title (Telegraf Bold 22pt) ─────────
    # Word maps non-bold Telegraf to UltraLight (invisible). Bold gives the visible weight.
    h1 = doc.styles["Heading 1"]
    h1.font.name = TELEGRAF
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = EVYE_BLACK
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.line_spacing = 1.15
    h1.paragraph_format.keep_with_next = True
    _set_font_xml(h1, TELEGRAF, FALLBACK_BODY)

    # ── Section Title (Telegraf UltraLight 20pt) ──────────
    h2 = doc.styles["Heading 2"]
    h2.font.name = TELEGRAF
    h2.font.size = Pt(20)
    h2.font.bold = False
    h2.font.color.rgb = EVYE_BLACK
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.keep_with_next = True
    _set_font_xml(h2, TELEGRAF, FALLBACK_BODY)

    # ── Subsection (Telegraf Bold 10pt) ───────────────────
    h3 = doc.styles["Heading 3"]
    h3.font.name = TELEGRAF
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = EVYE_BLACK
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    _set_font_xml(h3, TELEGRAF, FALLBACK_BODY)

    # ── Sub-subsection (Helvetica Neue Bold 9.5pt) ────────
    h4 = doc.styles["Heading 4"]
    h4.font.name = HELVETICA
    h4.font.size = Pt(9.5)
    h4.font.bold = True
    h4.font.color.rgb = EVYE_BLACK
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(3)
    h4.paragraph_format.keep_with_next = True
    _set_font_xml(h4, HELVETICA, FALLBACK_BODY)

    # ── Doc Subtitle (Helvetica Neue 8pt uppercase tracking) ──
    subtitle = doc.styles.add_style("Evye Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.base_style = doc.styles["Normal"]
    subtitle.font.name = HELVETICA
    subtitle.font.size = Pt(8)
    subtitle.font.bold = False
    subtitle.font.color.rgb = EVYE_BLACK
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    # Spacing/tracking set via XML (value in 1/20 of a point, so 30 = 1.5pt tracking)
    _set_font_xml(subtitle, HELVETICA, FALLBACK_BODY)
    _set_char_spacing(subtitle, 20)  # 1pt tracking — tight and refined for 8pt text

    # ── Cover Title (Telegraf UltraLight 28pt) ────────────
    cover = doc.styles.add_style("Evye Cover Title", WD_STYLE_TYPE.PARAGRAPH)
    cover.base_style = doc.styles["Normal"]
    cover.font.name = TELEGRAF
    cover.font.size = Pt(28)
    cover.font.bold = False
    cover.font.color.rgb = EVYE_BLACK
    cover.paragraph_format.space_before = Pt(140)
    cover.paragraph_format.space_after = Pt(8)
    cover.paragraph_format.line_spacing = 1.15
    _set_font_xml(cover, TELEGRAF, FALLBACK_BODY)


def _create_body_styles(doc):
    """Body text styles: justified body, left-aligned body, salutation, intro."""

    # ── Body Justified ────────────────────────────────────
    body_j = doc.styles.add_style("Evye Body", WD_STYLE_TYPE.PARAGRAPH)
    body_j.base_style = doc.styles["Normal"]
    body_j.font.name = HELVETICA
    body_j.font.size = Pt(9.5)
    body_j.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_j.paragraph_format.space_after = Pt(5)
    body_j.paragraph_format.line_spacing = 1.35
    _set_font_xml(body_j, HELVETICA, FALLBACK_BODY)

    # ── Body Left ─────────────────────────────────────────
    body_l = doc.styles.add_style("Evye Body Left", WD_STYLE_TYPE.PARAGRAPH)
    body_l.base_style = doc.styles["Normal"]
    body_l.font.name = HELVETICA
    body_l.font.size = Pt(9.5)
    body_l.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body_l.paragraph_format.space_after = Pt(5)
    body_l.paragraph_format.line_spacing = 1.35
    _set_font_xml(body_l, HELVETICA, FALLBACK_BODY)

    # ── Salutation ────────────────────────────────────────
    sal = doc.styles.add_style("Evye Salutation", WD_STYLE_TYPE.PARAGRAPH)
    sal.base_style = doc.styles["Normal"]
    sal.font.name = HELVETICA
    sal.font.size = Pt(9.5)
    sal.paragraph_format.space_before = Pt(0)
    sal.paragraph_format.space_after = Pt(8)
    _set_font_xml(sal, HELVETICA, FALLBACK_BODY)

    # ── Small text (8.5pt for descriptions, table notes) ──
    small = doc.styles.add_style("Evye Small", WD_STYLE_TYPE.PARAGRAPH)
    small.base_style = doc.styles["Normal"]
    small.font.name = HELVETICA
    small.font.size = Pt(8.5)
    small.font.color.rgb = EVYE_BLACK
    small.paragraph_format.space_after = Pt(3)
    _set_font_xml(small, HELVETICA, FALLBACK_BODY)

    # ── Tiny text (7pt for sidebar-style info, meta) ──────
    tiny = doc.styles.add_style("Evye Tiny", WD_STYLE_TYPE.PARAGRAPH)
    tiny.base_style = doc.styles["Normal"]
    tiny.font.name = HELVETICA
    tiny.font.size = Pt(7)
    tiny.font.color.rgb = EVYE_GREY
    tiny.paragraph_format.space_after = Pt(1)
    _set_font_xml(tiny, HELVETICA, FALLBACK_BODY)


def _create_clause_styles(doc):
    """Numbered clause styles for contracts."""

    # ── Clause heading (bold 10pt with number) ────────────
    clause = doc.styles.add_style("Evye Clause", WD_STYLE_TYPE.PARAGRAPH)
    clause.base_style = doc.styles["Normal"]
    clause.font.name = HELVETICA
    clause.font.size = Pt(10)
    clause.font.bold = True
    clause.paragraph_format.space_before = Pt(10)
    clause.paragraph_format.space_after = Pt(4)
    clause.paragraph_format.keep_with_next = True
    # Tab stop at 12mm for number/title alignment
    clause.paragraph_format.tab_stops.add_tab_stop(Mm(12), WD_TAB_ALIGNMENT.LEFT)
    _set_font_xml(clause, HELVETICA, FALLBACK_BODY)

    # ── Subclause (9.5pt with indent) ─────────────────────
    sub = doc.styles.add_style("Evye Subclause", WD_STYLE_TYPE.PARAGRAPH)
    sub.base_style = doc.styles["Normal"]
    sub.font.name = HELVETICA
    sub.font.size = Pt(9.5)
    sub.font.bold = False
    sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sub.paragraph_format.space_after = Pt(4)
    sub.paragraph_format.left_indent = Mm(12)
    sub.paragraph_format.first_line_indent = Mm(-9)
    sub.paragraph_format.tab_stops.add_tab_stop(Mm(12), WD_TAB_ALIGNMENT.LEFT)
    _set_font_xml(sub, HELVETICA, FALLBACK_BODY)

    # ── Sub-subclause (deeper indent) ─────────────────────
    subsub = doc.styles.add_style("Evye Sub-Subclause", WD_STYLE_TYPE.PARAGRAPH)
    subsub.base_style = doc.styles["Normal"]
    subsub.font.name = HELVETICA
    subsub.font.size = Pt(9.5)
    subsub.font.bold = False
    subsub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    subsub.paragraph_format.space_after = Pt(4)
    subsub.paragraph_format.left_indent = Mm(21)
    subsub.paragraph_format.first_line_indent = Mm(-9)
    subsub.paragraph_format.tab_stops.add_tab_stop(Mm(21), WD_TAB_ALIGNMENT.LEFT)
    _set_font_xml(subsub, HELVETICA, FALLBACK_BODY)


def _create_list_styles(doc):
    """Bullet and numbered list styles."""

    # ── Bullet Level 1 (disc) ─────────────────────────────
    bullet1 = doc.styles.add_style("Evye Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet1.base_style = doc.styles["Normal"]
    bullet1.font.name = HELVETICA
    bullet1.font.size = Pt(9.5)
    bullet1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet1.paragraph_format.space_after = Pt(2.5)
    bullet1.paragraph_format.left_indent = Mm(10)
    bullet1.paragraph_format.first_line_indent = Mm(-5)
    # Tab stop so bullet character and text align
    bullet1.paragraph_format.tab_stops.add_tab_stop(Mm(10), WD_TAB_ALIGNMENT.LEFT)
    _set_font_xml(bullet1, HELVETICA, FALLBACK_BODY)

    # ── Bullet Level 2 ───────────────────────────────────
    bullet2 = doc.styles.add_style("Evye Bullet 2", WD_STYLE_TYPE.PARAGRAPH)
    bullet2.base_style = doc.styles["Normal"]
    bullet2.font.name = HELVETICA
    bullet2.font.size = Pt(9.5)
    bullet2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet2.paragraph_format.space_after = Pt(2.5)
    bullet2.paragraph_format.left_indent = Mm(16)
    bullet2.paragraph_format.first_line_indent = Mm(-5)
    bullet2.paragraph_format.tab_stops.add_tab_stop(Mm(16), WD_TAB_ALIGNMENT.LEFT)
    _set_font_xml(bullet2, HELVETICA, FALLBACK_BODY)

    # ── Bullet Level 3 (dash) ────────────────────────────
    bullet3 = doc.styles.add_style("Evye Bullet 3", WD_STYLE_TYPE.PARAGRAPH)
    bullet3.base_style = doc.styles["Normal"]
    bullet3.font.name = HELVETICA
    bullet3.font.size = Pt(9.5)
    bullet3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet3.paragraph_format.space_after = Pt(2.5)
    bullet3.paragraph_format.left_indent = Mm(22)
    bullet3.paragraph_format.first_line_indent = Mm(-5)
    bullet3.paragraph_format.tab_stops.add_tab_stop(Mm(22), WD_TAB_ALIGNMENT.LEFT)
    _set_font_xml(bullet3, HELVETICA, FALLBACK_BODY)

    # ── Small bullet (8.5pt for fee table items) ──────────
    bullet_sm = doc.styles.add_style("Evye Bullet Small", WD_STYLE_TYPE.PARAGRAPH)
    bullet_sm.base_style = doc.styles["Normal"]
    bullet_sm.font.name = HELVETICA
    bullet_sm.font.size = Pt(8.5)
    bullet_sm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet_sm.paragraph_format.space_after = Pt(1.5)
    bullet_sm.paragraph_format.left_indent = Mm(5)
    bullet_sm.paragraph_format.first_line_indent = Mm(-3)
    _set_font_xml(bullet_sm, HELVETICA, FALLBACK_BODY)


def _create_table_styles(doc):
    """Create table styles for fee tables and key-value tables."""

    # ── Fee Table Header (bold 8pt) ───────────────────────
    th = doc.styles.add_style("Evye Table Header", WD_STYLE_TYPE.PARAGRAPH)
    th.base_style = doc.styles["Normal"]
    th.font.name = HELVETICA
    th.font.size = Pt(8)
    th.font.bold = True
    th.font.color.rgb = EVYE_BLACK
    th.paragraph_format.space_before = Pt(2)
    th.paragraph_format.space_after = Pt(2)
    th.paragraph_format.right_indent = Mm(0)  # Override Normal's right indent inside tables
    _set_font_xml(th, HELVETICA, FALLBACK_BODY)

    # ── Fee Table Body (8.5pt) ────────────────────────────
    tb = doc.styles.add_style("Evye Table Body", WD_STYLE_TYPE.PARAGRAPH)
    tb.base_style = doc.styles["Normal"]
    tb.font.name = HELVETICA
    tb.font.size = Pt(8.5)
    tb.font.color.rgb = EVYE_BLACK
    tb.paragraph_format.space_before = Pt(2)
    tb.paragraph_format.space_after = Pt(2)
    tb.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tb.paragraph_format.right_indent = Mm(0)
    _set_font_xml(tb, HELVETICA, FALLBACK_BODY)

    # ── Fee Table Item Title (bold 8.5pt) ─────────────────
    tit = doc.styles.add_style("Evye Table Item Title", WD_STYLE_TYPE.PARAGRAPH)
    tit.base_style = doc.styles["Normal"]
    tit.font.name = HELVETICA
    tit.font.size = Pt(8.5)
    tit.font.bold = True
    tit.font.color.rgb = EVYE_BLACK
    tit.paragraph_format.space_before = Pt(2)
    tit.paragraph_format.space_after = Pt(1)
    tit.paragraph_format.right_indent = Mm(0)
    _set_font_xml(tit, HELVETICA, FALLBACK_BODY)

    # ── Fee Table Total (bold, right-aligned) ─────────────
    tt = doc.styles.add_style("Evye Table Total", WD_STYLE_TYPE.PARAGRAPH)
    tt.base_style = doc.styles["Normal"]
    tt.font.name = HELVETICA
    tt.font.size = Pt(8.5)
    tt.font.bold = True
    tt.font.color.rgb = EVYE_BLACK
    tt.paragraph_format.space_before = Pt(2)
    tt.paragraph_format.space_after = Pt(2)
    tt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tt.paragraph_format.right_indent = Mm(0)
    _set_font_xml(tt, HELVETICA, FALLBACK_BODY)

    # ── KV Table Label (bold 9.5pt) ───────────────────────
    kvl = doc.styles.add_style("Evye KV Label", WD_STYLE_TYPE.PARAGRAPH)
    kvl.base_style = doc.styles["Normal"]
    kvl.font.name = HELVETICA
    kvl.font.size = Pt(9.5)
    kvl.font.bold = True
    kvl.paragraph_format.space_before = Pt(1)
    kvl.paragraph_format.space_after = Pt(1)
    kvl.paragraph_format.right_indent = Mm(0)
    _set_font_xml(kvl, HELVETICA, FALLBACK_BODY)

    # ── KV Table Value (9.5pt) ────────────────────────────
    kvv = doc.styles.add_style("Evye KV Value", WD_STYLE_TYPE.PARAGRAPH)
    kvv.base_style = doc.styles["Normal"]
    kvv.font.name = HELVETICA
    kvv.font.size = Pt(9.5)
    kvv.paragraph_format.space_before = Pt(1)
    kvv.paragraph_format.space_after = Pt(1)
    kvv.paragraph_format.right_indent = Mm(0)
    _set_font_xml(kvv, HELVETICA, FALLBACK_BODY)


def _create_special_styles(doc):
    """Signature blocks, acceptance, annex titles."""

    # ── Signature Name (bold 9.5pt) ───────────────────────
    sn = doc.styles.add_style("Evye Signatory Name", WD_STYLE_TYPE.PARAGRAPH)
    sn.base_style = doc.styles["Normal"]
    sn.font.name = HELVETICA
    sn.font.size = Pt(9.5)
    sn.font.bold = True
    sn.paragraph_format.space_before = Pt(0)
    sn.paragraph_format.space_after = Pt(0)
    _set_font_xml(sn, HELVETICA, FALLBACK_BODY)

    # ── Signature Title (9.5pt) ───────────────────────────
    st = doc.styles.add_style("Evye Signatory Title", WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles["Normal"]
    st.font.name = HELVETICA
    st.font.size = Pt(9.5)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)
    _set_font_xml(st, HELVETICA, FALLBACK_BODY)

    # ── Annex Title (Telegraf UltraLight 20pt) ────────────
    at = doc.styles.add_style("Evye Annex Title", WD_STYLE_TYPE.PARAGRAPH)
    at.base_style = doc.styles["Normal"]
    at.font.name = TELEGRAF
    at.font.size = Pt(20)
    at.font.bold = False
    at.font.color.rgb = EVYE_BLACK
    at.paragraph_format.space_before = Pt(0)
    at.paragraph_format.space_after = Pt(4)
    at.paragraph_format.page_break_before = True
    _set_font_xml(at, TELEGRAF, FALLBACK_BODY)


def _setup_header(doc, section):
    """Add logo and company info to the page header."""
    header = section.header
    header.is_linked_to_previous = False

    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()

    # Create a table for header layout: content area | logo + sidebar info
    # Full width: 210 - 20 (left) - 15 (right) = 175mm
    tbl = header.add_table(rows=1, cols=2, width=Mm(175))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove table borders
    _remove_table_borders(tbl)

    # Left cell: empty (content flows below)
    left_cell = tbl.cell(0, 0)
    left_cell.width = Mm(130)
    lp = left_cell.paragraphs[0]
    lp.text = ""

    # Right cell: logo + company info
    right_cell = tbl.cell(0, 1)
    right_cell.width = Mm(45)

    # Clear default paragraph and reset right indent (inherited from Normal)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.right_indent = Mm(0)

    # Also reset left cell right indent
    lp.paragraph_format.right_indent = Mm(0)

    # Add logo
    logo_path = ASSETS_DIR / "evye-logo.png"
    if logo_path.exists():
        run = rp.add_run()
        run.add_picture(str(logo_path), width=Mm(18))

    # Company info
    info_lines = [
        ("Evye LLP", True),
        ("298 Tiong Bahru Road", False),
        ("#05-01 Central Plaza", False),
        ("Singapore 168730", False),
        ("", False),
        ("www.evye.co", False),
    ]

    for text, is_bold in info_lines:
        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.right_indent = Mm(0)
        run = p.add_run(text)
        run.font.name = HELVETICA
        run.font.size = Pt(7)
        run.font.color.rgb = EVYE_BLACK
        run.bold = is_bold
        _set_run_font(run, HELVETICA, FALLBACK_BODY)


def _setup_footer(doc, section):
    """Page number (left) and confidentiality (right) in footer."""
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    # Create footer table: page number | confidentiality — spans full width
    tbl = footer.add_table(rows=1, cols=2, width=Mm(175))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(tbl)

    # Left: page number
    left_cell = tbl.cell(0, 0)
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.right_indent = Mm(0)
    run = lp.add_run("Page ")
    run.font.name = HELVETICA
    run.font.size = Pt(7)
    run.font.color.rgb = EVYE_BLACK
    _set_run_font(run, HELVETICA, FALLBACK_BODY)

    # Add PAGE field
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
        f'<w:r><w:rPr><w:rFonts w:ascii="{HELVETICA}" w:hAnsi="{HELVETICA}" w:cs="{FALLBACK_BODY}"/>'
        f'<w:sz w:val="14"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    lp._p.append(parse_xml(fld_xml))

    run2 = lp.add_run(" of ")
    run2.font.name = HELVETICA
    run2.font.size = Pt(7)
    run2.font.color.rgb = EVYE_BLACK
    _set_run_font(run2, HELVETICA, FALLBACK_BODY)

    # Add NUMPAGES field
    fld_xml2 = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES ">'
        f'<w:r><w:rPr><w:rFonts w:ascii="{HELVETICA}" w:hAnsi="{HELVETICA}" w:cs="{FALLBACK_BODY}"/>'
        f'<w:sz w:val="14"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
    )
    lp._p.append(parse_xml(fld_xml2))

    # Right: confidentiality
    right_cell = tbl.cell(0, 1)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.line_spacing = 1.5
    rp.paragraph_format.right_indent = Mm(0)

    conf_lines = [
        "Private & Confidential",
        "\u00A9 Evye LLP",
        "All Rights Reserved",
    ]
    for i, line in enumerate(conf_lines):
        if i > 0:
            rp.add_run("\n")
        run = rp.add_run(line)
        run.font.name = HELVETICA
        run.font.size = Pt(6.5)
        run.font.color.rgb = EVYE_GREY
        _set_run_font(run, HELVETICA, FALLBACK_BODY)


def _add_sample_content(doc):
    """Add sample content demonstrating every style."""

    # Remove the default empty paragraph
    if doc.paragraphs:
        _delete_paragraph(doc.paragraphs[0])

    # Document Title
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run("Document Title")
    run.font.name = TELEGRAF
    run.font.size = Pt(22)
    run.font.color.rgb = EVYE_BLACK
    run.bold = True
    _set_run_font(run, TELEGRAF, FALLBACK_BODY)

    # Subtitle
    p = doc.add_paragraph(style="Evye Subtitle")
    p.add_run("SUBTITLE TEXT IN UPPERCASE WITH TRACKING")

    # Salutation
    p = doc.add_paragraph(style="Evye Salutation")
    p.add_run("Dear Mr Doe,")

    # Body text
    p = doc.add_paragraph(style="Evye Body")
    p.add_run(
        "We are pleased to present this document as a demonstration of the Evye master template styles. "
        "This paragraph uses the Evye Body style with justified alignment, Helvetica Neue at 9.5pt, "
        "and 1.35x line spacing. All body text in Evye documents should use this style."
    )

    # ── Heading hierarchy ─────────────────────────────────
    doc.add_paragraph("Section Title", style="Heading 2")

    p = doc.add_paragraph(style="Evye Body")
    p.add_run("This section demonstrates the heading hierarchy. Heading 2 uses Telegraf UltraLight at 20pt.")

    doc.add_paragraph("Subsection Title", style="Heading 3")

    p = doc.add_paragraph(style="Evye Body")
    p.add_run("Heading 3 uses Telegraf Bold at 10pt for subsection titles.")

    doc.add_paragraph("Sub-subsection Title", style="Heading 4")

    p = doc.add_paragraph(style="Evye Body")
    p.add_run("Heading 4 uses Helvetica Neue Bold at 9.5pt for the finest heading level.")

    # ── Clause numbering ──────────────────────────────────
    doc.add_paragraph("Clause Numbering", style="Heading 2")

    p = doc.add_paragraph(style="Evye Clause")
    p.add_run("1.\tDEFINITIONS AND INTERPRETATION")

    p = doc.add_paragraph(style="Evye Subclause")
    p.add_run('1.1.\tIn this Agreement, unless the context otherwise requires, the following terms shall have the meanings ascribed to them below.')

    p = doc.add_paragraph(style="Evye Subclause")
    p.add_run('1.2.\t"Client" means the party engaging Evye LLP for the provision of services as described in this Agreement.')

    p = doc.add_paragraph(style="Evye Sub-Subclause")
    p.add_run('1.2.1.\tThis includes any subsidiary, affiliate, or related entity of the Client as may be notified in writing.')

    p = doc.add_paragraph(style="Evye Clause")
    p.add_run("2.\tSCOPE OF SERVICES")

    p = doc.add_paragraph(style="Evye Subclause")
    p.add_run("2.1.\tEvye shall provide the services as described in the Annexes attached to this Agreement.")

    # ── Bullet lists ──────────────────────────────────────
    doc.add_paragraph("Bullet Lists", style="Heading 2")

    p = doc.add_paragraph(style="Evye Body Left")
    p.add_run("The following items are included in the scope of work:")

    bullets = [
        "Brand strategy and positioning workshop",
        "Visual identity design including logo, colour palette, and typography",
        "Brand guidelines document with usage rules",
        "Stationery suite: business cards, letterhead, envelope",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="Evye Bullet")
        p.add_run(f"\u2022\t{b}")

    # Level 2 bullets
    p = doc.add_paragraph(style="Evye Body Left")
    p.add_run("With sub-items:")

    p = doc.add_paragraph(style="Evye Bullet")
    p.add_run("\u2022\tPrimary deliverables")

    sub_bullets = ["Logo in vector and raster formats", "Colour palette with Pantone, CMYK, and HEX values"]
    for b in sub_bullets:
        p = doc.add_paragraph(style="Evye Bullet 2")
        p.add_run(f"\u2022\t{b}")

    # Level 3 bullets (dash)
    dash_items = ["Print-ready files", "Digital-optimised files"]
    for d in dash_items:
        p = doc.add_paragraph(style="Evye Bullet 3")
        p.add_run(f"\u2013\t{d}")

    # ── Fee table ─────────────────────────────────────────
    doc.add_paragraph("Fee Table", style="Heading 2")

    p = doc.add_paragraph(style="Evye Small")
    p.add_run("Below is a sample fee schedule for the project:")

    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    _set_cell(hdr.cells[0], "Item", "Evye Table Header", doc)
    _set_cell(hdr.cells[1], "Description", "Evye Table Header", doc)
    _set_cell(hdr.cells[2], "Fee", "Evye Table Header", doc)

    # Data rows
    items = [
        ("Brand Strategy", "Discovery, research, and positioning workshop", "$ 4,500"),
        ("Visual Identity", "Logo design, colour palette, typography system", "$ 8,000"),
        ("Brand Guidelines", "Comprehensive usage document (40-60 pages)", "$ 3,500"),
    ]
    for i, (name, desc, fee) in enumerate(items):
        row = table.rows[i + 1]
        # Item name (bold)
        _set_cell(row.cells[0], name, "Evye Table Item Title", doc)
        _set_cell(row.cells[1], desc, "Evye Table Body", doc)
        _set_cell(row.cells[2], fee, "Evye Table Body", doc)

    # Total row
    total_row = table.rows[4]
    _set_cell(total_row.cells[0], "", "Evye Table Body", doc)
    _set_cell(total_row.cells[1], "Total", "Evye Table Total", doc)
    _set_cell(total_row.cells[2], "$ 16,000", "Evye Table Total", doc)

    # Apply borders
    _apply_fee_table_borders(table)

    # Set column widths (content area ~128mm after right indent)
    for row in table.rows:
        row.cells[0].width = Mm(28)
        row.cells[1].width = Mm(72)
        row.cells[2].width = Mm(28)

    # ── Key-value table ───────────────────────────────────
    doc.add_paragraph("Key-Value Details", style="Heading 2")

    kv_table = doc.add_table(rows=5, cols=2)
    kv_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(kv_table)

    kv_data = [
        ("Role Title:", "Brand Strategist"),
        ("Location:", "Singapore (Remote)"),
        ("Job Type:", "Independent Contractor"),
        ("Start Date:", "1 May 2026"),
        ("Reporting To:", "Managing Partner"),
    ]
    for i, (label, value) in enumerate(kv_data):
        _set_cell(kv_table.rows[i].cells[0], label, "Evye KV Label", doc)
        _set_cell(kv_table.rows[i].cells[1], value, "Evye KV Value", doc)
        kv_table.rows[i].cells[0].width = Mm(35)
        kv_table.rows[i].cells[1].width = Mm(90)

    # ── Signature block ───────────────────────────────────
    doc.add_paragraph("Signature Block", style="Heading 2")

    # Signature space
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph(style="Evye Signatory Name")
    p.add_run("Michael Ryan Chan")

    p = doc.add_paragraph(style="Evye Signatory Title")
    p.add_run("Managing Partner")

    p = doc.add_paragraph(style="Evye Signatory Title")
    p.add_run("Evye LLP")


# ── Helper functions ─────────────────────────────────────────

def _set_font_xml(style, ascii_font, cs_font):
    """Set font names at the XML level for full compatibility."""
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is not None:
        rpr.remove(rFonts)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{ascii_font}" w:hAnsi="{ascii_font}" '
        f'w:cs="{cs_font}" w:eastAsia="{ascii_font}"/>'
    )
    rpr.insert(0, rFonts)


def _set_run_font(run, ascii_font, cs_font):
    """Set font names on a specific run at the XML level."""
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._r.insert(0, rpr)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is not None:
        rpr.remove(rFonts)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{ascii_font}" w:hAnsi="{ascii_font}" '
        f'w:cs="{cs_font}" w:eastAsia="{ascii_font}"/>'
    )
    rpr.insert(0, rFonts)


def _set_char_spacing(style, spacing):
    """Set character spacing (tracking) via XML. Value is in twips (1/20 of a point)."""
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    spacing_xml = parse_xml(
        f'<w:spacing {nsdecls("w")} w:val="{int(spacing)}"/>'
    )
    rpr.append(spacing_xml)


def _set_font_weight(style, bold):
    """Explicitly set bold on/off in XML to control font weight rendering."""
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    # Remove existing bold element
    for b in rpr.findall(qn("w:b")):
        rpr.remove(b)
    val = "true" if bold else "false"
    b_xml = parse_xml(f'<w:b {nsdecls("w")} w:val="{val}"/>')
    rpr.append(b_xml)


def _delete_paragraph(paragraph):
    """Remove a paragraph from the document body."""
    p = paragraph._p
    p.getparent().remove(p)


def _set_cell(cell, text, style_name, doc):
    """Set cell text with a named style."""
    p = cell.paragraphs[0]
    p.text = ""
    p.style = doc.styles[style_name]
    if text:
        run = p.add_run(text)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _apply_fee_table_borders(table):
    """Apply brand borders to a fee table: top/bottom on header, light grey between rows."""
    tbl = table._tbl
    # Table-level: no borders (we set per-cell)
    _remove_table_borders(table)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)

            if i == 0:
                # Header row: top and bottom black borders
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'</w:tcBorders>'
                )
            elif i == len(table.rows) - 1:
                # Total row: top double black border
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                    f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'</w:tcBorders>'
                )
            else:
                # Data rows: light grey bottom border
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                    f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'</w:tcBorders>'
                )

            existing = tcPr.find(qn("w:tcBorders"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(borders)


if __name__ == "__main__":
    create_master_template()
