"""
Evye LLP — Word Document Generator
====================================
Generates branded .docx files by cloning evye-master.docx and populating
it with structured content. The master template's header (logo, address,
date auto-field) and footer (page numbers, confidentiality) are preserved
automatically on every generated document.

Usage
-----
    from word_generator import generate_docx

    path = generate_docx("letter", {
        "doc_title": "Notice of Engagement",
        "doc_subtitle": "INDEPENDENT CONTRACTOR",
        "recipient_salutation": "Mr",
        "recipient_name": "John Doe",
        "body": [
            {"type": "paragraph", "text": "We are pleased to..."},
            {"type": "heading2", "text": "Scope of Work"},
            {"type": "bullet", "text": "Brand strategy workshop"},
        ],
        "signatory_name": "Michael Ryan Chan",
        "signatory_title": "Managing Partner",
    })
    # → "output/2026-04-11_letter_notice-of-engagement.docx"

Document Types
--------------
  "letter"     — General correspondence, memos, notices
  "contract"   — Contractor/employment agreements with numbered clauses
  "quotation"  — Project proposals with fee tables and payment milestones
  "custom"     — Free-form: supply a `sections` list and it builds what you give it

Content Block Types (used in body/sections lists)
--------------------------------------------------
  {"type": "paragraph",   "text": "..."}         — Body style (justified)
  {"type": "paragraph_left", "text": "..."}      — Body Left style
  {"type": "heading2",    "text": "..."}         — H2 section title (14pt Telegraf Bold)
  {"type": "heading3",    "text": "..."}         — H3 subsection (10pt UPPERCASE)
  {"type": "heading4",    "text": "..."}         — H4 sub-subsection (10pt Bold Italic)
  {"type": "bullet",      "text": "...", "level": 1|2|3}  — bullet list
  {"type": "clause",      "number": "1.", "title": "SCOPE", "text": "..."}
  {"type": "subclause",   "number": "1.1.", "text": "..."}
  {"type": "sub_subclause","number": "1.1.1.", "text": "..."}
  {"type": "fee_table",   "rows": [{"item":"...","description":"...","fee":"$X,XXX"}], "total": "$X,XXX"}
  {"type": "kv_table",    "rows": [{"label":"Role Title:","value":"Brand Strategist"}]}
  {"type": "spacer"}                             — empty paragraph
  {"type": "page_break"}                         — explicit page break
"""

import shutil
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_ROOT = Path(__file__).resolve().parent
MASTER_PATH  = PROJECT_ROOT / "templates" / "evye-master.docx"
OUTPUT_DIR   = PROJECT_ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

__version__ = "1.2.0-dev"   # bumped to 1.2.0 at release (Task 12)


# ── Public API ────────────────────────────────────────────────────────────────

def generate_docx(doc_type: str, data: dict, filename: str = None) -> str:
    """
    Generate a branded .docx file.

    Parameters
    ----------
    doc_type : str
        One of: "letter", "contract", "quotation", "custom"
    data : dict
        Content data — see module docstring for available keys.
    filename : str, optional
        Override the output filename (without extension).

    Returns
    -------
    str
        Absolute path to the generated .docx file.
    """
    # Clone the master template
    doc = Document(str(MASTER_PATH))

    # Clear all existing body paragraphs and tables
    _clear_body(doc)

    # Route to builder
    builders = {
        "letter":    _build_letter,
        "contract":  _build_contract,
        "quotation": _build_quotation,
        "custom":    _build_custom,
    }
    builder = builders.get(doc_type, _build_custom)
    builder(doc, data)

    # Determine output path
    if not filename:
        today = date.today().strftime("%Y-%m-%d")
        slug = _slugify(data.get("doc_title", doc_type))
        filename = f"{today}_{doc_type}_{slug}"
    output_path = OUTPUT_DIR / f"{filename}.docx"
    # Avoid overwriting an existing file with the same auto-slug
    if output_path.exists():
        counter = 2
        while (OUTPUT_DIR / f"{filename}-{counter}.docx").exists():
            counter += 1
        output_path = OUTPUT_DIR / f"{filename}-{counter}.docx"

    # Strip the <w:embedTrueTypeFonts/> flag that causes Word to bundle ~12 MB
    # of system fonts (Calibri, Arial, Times, Courier, Cambria) into every
    # document. Those fonts are universally pre-installed; embedding them
    # wastes space without benefit. Telegraf (our custom font) is delivered
    # separately via the /fonts bot command — so nothing changes for rendering.
    from docx.oxml.ns import qn as _qn
    settings_elem = doc.settings.element
    for _elem in settings_elem.findall(_qn('w:embedTrueTypeFonts')):
        settings_elem.remove(_elem)

    doc.save(str(output_path))
    return str(output_path)


# ── Document builders ─────────────────────────────────────────────────────────

def _build_letter(doc, data):
    """General letter / correspondence."""
    _add_title_block(doc, data)
    if data.get("recipient_salutation") or data.get("recipient_name"):
        sal = f"Dear {data.get('recipient_salutation','')} {data.get('recipient_name','')},"
        _add_para(doc, sal.strip(), "Salutation")
        _add_spacer(doc)
    _render_blocks(doc, data.get("body", []))
    _add_spacer(doc)
    _add_signature(doc, data)


def _build_contract(doc, data):
    """Contractor / employment agreement with numbered clauses."""
    _add_title_block(doc, data)

    # Key-value details table (role info)
    kv_rows = data.get("kv_details", [])
    if kv_rows:
        _add_kv_table(doc, kv_rows)
        _add_spacer(doc)

    if data.get("recipient_salutation") or data.get("recipient_name"):
        sal = f"Dear {data.get('recipient_salutation','')} {data.get('recipient_name','')},"
        _add_para(doc, sal.strip(), "Salutation")
        _add_spacer(doc)

    if data.get("intro"):
        _add_para(doc, data["intro"], "Body")
        _add_spacer(doc)

    # Clauses
    for clause in data.get("clauses", []):
        _add_clause(doc, clause)

    _add_spacer(doc)
    _add_signature(doc, data)

    # Annexes
    for annex in data.get("annexes", []):
        _add_page_break(doc)
        _add_para(doc, annex.get("title", "Annex"), "Annex Title")
        if annex.get("subtitle"):
            _add_para(doc, annex["subtitle"], "Heading 3")
        _render_blocks(doc, annex.get("body", []))


def _build_quotation(doc, data):
    """Project quotation / proposal with fee tables."""
    _add_title_block(doc, data)

    if data.get("intro"):
        _add_para(doc, data["intro"], "Body")
        _add_spacer(doc)

    for phase in data.get("phases", []):
        _add_para(doc, phase.get("title", ""), "Heading 2")
        if phase.get("description"):
            _add_para(doc, phase["description"], "Body")
        if phase.get("line_items"):
            fee_rows = [
                {"item": li.get("name",""), "description": li.get("description",""), "fee": li.get("fee","")}
                for li in phase["line_items"]
            ]
            _add_fee_table(doc, fee_rows, subtotal=phase.get("subtotal"), total=phase.get("total"))
        _add_spacer(doc)

    if data.get("cost_summary"):
        cs = data["cost_summary"]
        _add_para(doc, cs.get("title", "Cost Summary"), "Heading 2")
        rows = [{"item": li.get("name",""), "description": "", "fee": li.get("fee","")}
                for li in cs.get("line_items", [])]
        _add_fee_table(doc, rows, total=cs.get("total"))
        _add_spacer(doc)

    if data.get("payment_milestones"):
        _add_para(doc, "Payment Schedule", "Heading 2")
        for i, m in enumerate(data["payment_milestones"], 1):
            _add_para(doc, f"{i}.\t{m.get('amount','')} — {m.get('description','')}", "Subclause")
        _add_spacer(doc)

    _add_signature(doc, data)


def _build_custom(doc, data):
    """Free-form: supply title block + a list of content blocks."""
    _add_title_block(doc, data)
    _render_blocks(doc, data.get("body", []))
    if data.get("signatory_name") or data.get("signatory_title"):
        _add_spacer(doc)
        _add_signature(doc, data)


# ── Rendering helpers ─────────────────────────────────────────────────────────

def _render_blocks(doc, blocks):
    _assign_numbers(blocks)
    for block in blocks:
        _render_block(doc, block)


def _render_block(doc, block):
    t = block.get("type", "paragraph")

    if t == "paragraph":
        _add_para(doc, block.get("runs") or block["text"], "Body")
    elif t == "paragraph_left":
        _add_para(doc, block.get("runs") or block["text"], "Body Left")
    elif t == "heading2":
        _add_para(doc, block["text"], "Heading 2")
    elif t == "heading3":
        _add_para(doc, block["text"], "Heading 3")
    elif t == "heading4":
        _add_para(doc, block["text"], "Heading 4")
    elif t == "bullet":
        level = block.get("level", 1)
        style = {1: "Bullet", 2: "Bullet 2", 3: "Bullet 3"}.get(level, "Bullet")
        # Evye Bullet uses a literal bullet + tab (not Word list numbering)
        content = block.get("runs") or block["text"]
        if isinstance(content, str) and content.startswith("•"):
            _add_para(doc, content, style)          # legacy pre-glyphed text
        else:
            _add_para(doc, content, style, prefix="•\t")
    elif t == "clause":
        _add_clause(doc, block)
    elif t == "subclause":
        text = f"{block.get('number','')}".rstrip() + "\t" + block.get("text","")
        _add_para(doc, text, "Subclause")
    elif t == "sub_subclause":
        text = f"{block.get('number','')}".rstrip() + "\t" + block.get("text","")
        _add_para(doc, text, "Sub-Subclause")
    elif t == "fee_table":
        _add_fee_table(doc, block.get("rows",[]), total=block.get("total"), headers=block.get("headers"))
    elif t == "kv_table":
        _add_kv_table(doc, block.get("rows",[]))
    elif t == "process_table":
        _add_process_table(doc, block.get("rows",[]), headers=block.get("headers"))
    elif t == "spacer":
        _add_spacer(doc)
    elif t == "page_break":
        _add_page_break(doc)
    elif t == "numbered":
        level = block.get("level", 1)
        style = {1: "Bullet", 2: "Bullet 2", 3: "Bullet 3"}.get(level, "Bullet")
        content = block.get("runs") or block.get("text", "")
        _add_para(doc, content, style, prefix=(block.get("number", "1.") + "\t"))
    elif t == "todo":
        level = block.get("level", 1)
        style = {1: "Bullet", 2: "Bullet 2", 3: "Bullet 3"}.get(level, "Bullet")
        glyph = "☑" if block.get("checked") else "☐"
        content = block.get("runs") or block.get("text", "")
        _add_para(doc, content, style, prefix=glyph + "\t")
    elif t == "code_block":
        _add_code_block(doc, block.get("text", ""))
    elif t == "image":
        _add_image(doc, block.get("url", ""), caption=block.get("caption"))
    elif t == "generic_table":
        _add_generic_table(doc, block.get("headers"), block.get("rows", []))


def _add_title_block(doc, data):
    if data.get("doc_title"):
        _add_para(doc, data["doc_title"], "Heading 1")
    if data.get("doc_subtitle"):
        _add_para(doc, data["doc_subtitle"], "Doc Subtitle")
    if data.get("doc_title") or data.get("doc_subtitle"):
        _add_spacer(doc)


def _add_clause(doc, clause):
    """Add a numbered clause with optional subclauses."""
    number = clause.get("number", "")
    title  = clause.get("title", "")
    text   = clause.get("text", "")

    # Clause heading: "1.\tSCOPE OF SERVICES"
    heading = f"{number}\t{title}".strip() if title else f"{number}\t{text}".strip()
    _add_para(doc, heading, "Clause")

    # If there's separate body text
    if title and text:
        _add_para(doc, f"{number[:-1]}.\t{text}", "Subclause") if number else _add_para(doc, text, "Body")

    for sub in clause.get("subclauses", []):
        sub_num  = sub.get("number", "")
        sub_text = sub.get("text", "")
        _add_para(doc, f"{sub_num}\t{sub_text}", "Subclause")
        for subsub in sub.get("subclauses", []):
            ss_num  = subsub.get("number", "")
            ss_text = subsub.get("text", "")
            _add_para(doc, f"{ss_num}\t{ss_text}", "Sub-Subclause")


def _add_signature(doc, data):
    import re
    _add_spacer(doc)
    _add_para(doc, data.get("signatory_name", "Michael Ryan Chan"), "Signatory Name")
    # Strip a trailing "Evye LLP" / "Evye Pte Ltd" from the title to avoid
    # duplication — the next line always renders "Evye LLP" explicitly.
    title = data.get("signatory_title", "Managing Partner")
    title = re.sub(r',?\s*Evye\s*(LLP|Pte\.?\s*Ltd|Ltd)?\.?\s*$', '', title, flags=re.I).strip()
    _add_para(doc, title or "Managing Partner", "Signatory Title")
    _add_para(doc, "Evye LLP", "Signatory Title")


def _assign_numbers(blocks):
    """
    Assign display numbers to `numbered` blocks in-place. Contiguous runs at
    the same level count up; deeper levels restart; any non-list block resets
    all counters; a bullet/todo at level L resets counters at levels >= L
    (so nested bullets under a numbered item don't break the parent count).
    """
    counters = {}
    for b in blocks:
        t = b.get("type")
        if t == "numbered":
            lvl = b.get("level", 1)
            counters[lvl] = counters.get(lvl, 0) + 1
            for k in [k for k in counters if k > lvl]:
                counters.pop(k)
            b["number"] = f"{counters[lvl]}."
        elif t in ("bullet", "todo"):
            lvl = b.get("level", 1)
            for k in [k for k in counters if k >= lvl]:
                counters.pop(k)
        else:
            counters = {}
    return blocks


def _add_code_block(doc, text):
    """Monospace block: Body Left style, Courier New 8pt, real line breaks."""
    p = doc.add_paragraph(style="Body Left")
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        if idx < len(lines) - 1:
            run.add_break()
    return p


def _add_image(doc, url, caption=None):
    """
    Download and embed an image, capped at 133mm width (the text column).
    Any failure degrades to a visible placeholder — never silent.
    """
    import urllib.request, tempfile, os
    from docx.shared import Emu
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "evye-docgen"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = resp.read()
        tmp = tempfile.NamedTemporaryFile(suffix=".img", delete=False)
        try:
            tmp.write(data)
            tmp.close()
            pic = doc.add_picture(tmp.name)
            max_w = Mm(133)
            if pic.width > max_w:
                pic.height = Emu(int(pic.height * (max_w / pic.width)))
                pic.width = max_w
        finally:
            os.unlink(tmp.name)
    except Exception:
        _add_para(doc, f"[Image unavailable: {url[:120]}]", "Small")
        return
    if caption:
        _add_para(doc, caption, "Small")


def _add_generic_table(doc, headers, rows):
    """
    N-column fallback for tables that don't fit kv/fee/process shapes
    (1 column, or 4+ columns). Equal widths totalling ~135mm, standard
    hairline styling. Preserves EVERY cell — no truncation.
    """
    n_cols = 0
    if rows:
        n_cols = max(len(r) for r in rows)
    if headers:
        n_cols = max(n_cols, len(headers))
    if n_cols == 0:
        return
    n_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    _set_table_no_borders(table)
    col_mm = round(135.0 / n_cols, 1)
    _set_table_fixed_layout(table, [col_mm] * n_cols)
    for i in range(n_cols):
        for cell in table.columns[i].cells:
            cell.width = Mm(col_mm)
    HAIRLINE = ("single", "2", "DDDDDD")
    offset = 0
    if headers:
        hdr = table.rows[0]
        padded = list(headers) + [""] * (n_cols - len(headers))
        for ci, h in enumerate(padded):
            _set_cell_style(hdr.cells[ci], "Table Header", h, WD_ALIGN_PARAGRAPH.LEFT,
                            fill="FFFFFF",
                            top=("single", "4", "000000"),
                            bottom=("single", "4", "000000"))
        offset = 1
    for ri, row in enumerate(rows):
        padded = list(row) + [""] * (n_cols - len(row))
        for ci, val in enumerate(padded):
            _set_cell_style(table.rows[ri + offset].cells[ci], "Table Body", val,
                            WD_ALIGN_PARAGRAPH.LEFT, fill="FFFFFF", bottom=HAIRLINE)


# ── Table builders ────────────────────────────────────────────────────────────

def _add_fee_table(doc, rows, subtotal=None, total=None, headers=None):
    """
    Add a styled fee table matching the master template's Fee Table design:
    - No grid borders, hairline row separators
    - 3 columns: Item | Description | Fee
    - Header row: top + bottom black rule
    - Data rows: light grey hairline bottom
    - Total row: black top rule
    """
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    table = doc.add_table(rows=1 + len(rows) + (1 if total else 0), cols=3)
    table.style = "Table Grid"  # will be overridden by cell-level XML

    # Remove table-level borders
    _set_table_no_borders(table)

    # Set column widths: Item 32mm, Description 80mm, Fee 23mm = 135mm
    # 133mm effective text column + 1.5mm right cell margin (85 dxa) = 134.5mm → 135mm
    _set_table_fixed_layout(table, [32, 80, 23])
    for i, width in enumerate([Mm(32), Mm(80), Mm(23)]):
        for cell in table.columns[i].cells:
            cell.width = width

    # Header row — use source headers if provided, else defaults
    if headers and len(headers) >= 3:
        h0, h1, h2 = headers[0], headers[1], headers[2]
    else:
        h0, h1, h2 = "Item", "Description", "Fee"
    hdr = table.rows[0]
    for ci, (text, align) in enumerate([(h0, WD_ALIGN_PARAGRAPH.LEFT),
                                          (h1, WD_ALIGN_PARAGRAPH.LEFT),
                                          (h2, WD_ALIGN_PARAGRAPH.RIGHT)]):
        cell = hdr.cells[ci]
        _set_cell_style(cell, "Table Header", text, align,
                        fill="FFFFFF",
                        top=("single","4","000000"),
                        bottom=("single","4","000000"))

    # Data rows
    HAIRLINE = ("single", "2", "DDDDDD")
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        _set_cell_style(row.cells[0], "Table Body",       row_data.get("item",""),        WD_ALIGN_PARAGRAPH.LEFT,  fill="FFFFFF", bottom=HAIRLINE)
        _set_cell_style(row.cells[1], "Table Body",       row_data.get("description",""), WD_ALIGN_PARAGRAPH.LEFT,  fill="FFFFFF", bottom=HAIRLINE)
        _set_cell_style(row.cells[2], "Table Body",       _fmt_fee(row_data.get("fee","")), WD_ALIGN_PARAGRAPH.RIGHT, fill="FFFFFF", bottom=HAIRLINE)

    # Total row
    if total:
        total_row = table.rows[-1]
        _set_cell_style(total_row.cells[0], "Table Total", "",             WD_ALIGN_PARAGRAPH.LEFT,  fill="FFFFFF", top=("single","4","000000"))
        _set_cell_style(total_row.cells[1], "Table Total", "Total",        WD_ALIGN_PARAGRAPH.RIGHT, fill="FFFFFF", top=("single","4","000000"))
        _set_cell_style(total_row.cells[2], "Table Total", _fmt_fee(total), WD_ALIGN_PARAGRAPH.RIGHT, fill="FFFFFF", top=("single","4","000000"))


def _add_kv_table(doc, rows):
    """
    Add a styled key-value table matching the master template's Key-Value Details design:
    - No grid borders, hairline row separators
    - 2 columns: Label | Value
    - First row: top black rule
    - All rows: hairline bottom border
    """
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"

    _set_table_no_borders(table)
    _set_table_fixed_layout(table, [42, 93])

    # Column widths: Label 42mm, Value 93mm = 135mm (133mm + 1.5mm right cell margin)
    for cell in table.columns[0].cells:
        cell.width = Mm(42)
    for cell in table.columns[1].cells:
        cell.width = Mm(93)

    HAIRLINE = ("single", "2", "DDDDDD")
    TOP_RULE = ("single", "4", "000000")

    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        top_border = TOP_RULE if ri == 0 else "none"
        _set_cell_style(row.cells[0], "KV Label", row_data.get("label",""),
                        WD_ALIGN_PARAGRAPH.LEFT, fill="FFFFFF",
                        top=top_border, bottom=HAIRLINE)
        _set_cell_style(row.cells[1], "KV Value", row_data.get("value",""),
                        WD_ALIGN_PARAGRAPH.LEFT, fill="FFFFFF",
                        top=top_border, bottom=HAIRLINE)


def _add_process_table(doc, rows, headers=None):
    """
    Add a 3-column process/step table (left-aligned, no fee column).
    Column widths: 20mm | 45mm | 80mm
    """
    n_rows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=n_rows, cols=3)
    table.style = "Table Grid"
    _set_table_no_borders(table)
    _set_table_fixed_layout(table, [18, 42, 75])

    # Column widths: Step 18mm, Stage 42mm, Description 75mm = 135mm (133mm + 1.5mm right cell margin)
    for i, width in enumerate([Mm(18), Mm(42), Mm(75)]):
        for cell in table.columns[i].cells:
            cell.width = width

    HAIRLINE = ("single", "2", "DDDDDD")
    row_offset = 0

    if headers:
        hdr = table.rows[0]
        for ci, h in enumerate(headers[:3]):
            _set_cell_style(hdr.cells[ci], "Table Header", h, WD_ALIGN_PARAGRAPH.LEFT,
                            fill="FFFFFF",
                            top=("single", "4", "000000"),
                            bottom=("single", "4", "000000"))
        row_offset = 1

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + row_offset]
        _set_cell_style(row.cells[0], "Table Body", row_data.get("col1", ""), WD_ALIGN_PARAGRAPH.LEFT, fill="FFFFFF", bottom=HAIRLINE)
        _set_cell_style(row.cells[1], "Table Body", row_data.get("col2", ""), WD_ALIGN_PARAGRAPH.LEFT, fill="FFFFFF", bottom=HAIRLINE)
        _set_cell_style(row.cells[2], "Table Body", row_data.get("col3", ""), WD_ALIGN_PARAGRAPH.LEFT, fill="FFFFFF", bottom=HAIRLINE)


# ── Low-level OOXML helpers ───────────────────────────────────────────────────

def _runs_plain(runs):
    """Plain-text concatenation of a run list."""
    return "".join(r.get("t", "") for r in (runs or [])).strip()


def _apply_run_format(run, fmt):
    if fmt.get("b"):
        run.bold = True
    if fmt.get("i"):
        run.italic = True
    if fmt.get("u"):
        run.underline = True
    if fmt.get("s"):
        run.font.strike = True
    if fmt.get("c"):
        run.font.name = "Courier New"


def _add_hyperlink_run(paragraph, text, url, fmt=None):
    """Append a clickable hyperlink run (underlined, Word-blue)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    fmt = fmt or {}
    props = ['<w:color w:val="1F4E79"/>', '<w:u w:val="single"/>']
    if fmt.get("b"):
        props.append('<w:b/>')
    if fmt.get("i"):
        props.append('<w:i/>')
    hyperlink_xml = (
        f'<w:hyperlink {nsdecls("w", "r")} r:id="{r_id}">'
        f'<w:r><w:rPr>{"".join(props)}</w:rPr>'
        f'<w:t xml:space="preserve">{_xml_escape(text)}</w:t></w:r>'
        f'</w:hyperlink>'
    )
    paragraph._p.append(parse_xml(hyperlink_xml))


def _add_para(doc, content, style_name, prefix=None):
    """
    Add a paragraph. `content` is a plain string OR a list of run dicts
    {"t","b","i","u","s","c","link"}. `prefix` ("•\t", "1.\t", "☐\t") is
    prepended as an unformatted run so list glyphs stay unstyled.
    """
    p = doc.add_paragraph(style=style_name)
    if prefix:
        p.add_run(prefix)
    if isinstance(content, str):
        p.add_run(content)
        return p
    for r in content:
        text = r.get("t", "")
        if not text:
            continue
        if r.get("link"):
            _add_hyperlink_run(p, text, r["link"], r)
        else:
            run = p.add_run(text)
            _apply_run_format(run, r)
    return p


def _add_spacer(doc):
    doc.add_paragraph(style="Body")


def _add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph(style="Normal")
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _clear_body(doc):
    """Remove all paragraphs and tables from the document body."""
    from docx.oxml.ns import qn
    body = doc.element.body
    # Remove everything except sectPr (section properties, holds header/footer links)
    to_remove = []
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != "sectPr":
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)
    # Add a single empty paragraph before sectPr (Word requires at least one)
    from lxml import etree
    empty_p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="Normal"/></w:pPr></w:p>')
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.insert(list(body).index(sect_pr), empty_p)
    else:
        body.append(empty_p)


def _set_table_fixed_layout(table, col_widths_mm):
    """
    Lock column widths by writing <w:tblGrid>, <w:tblLayout type="fixed">,
    and <w:tblW>. python-docx's cell.width alone is ignored because Word
    uses tblGrid as authoritative.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    # dxa = twentieths of a point; 1 mm ≈ 56.693 dxa
    dxa = [int(w * 56.693) for w in col_widths_mm]
    total = sum(dxa)

    # tblW — total width
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:type="dxa" w:w="{total}"/>'))

    # tblLayout — fixed (don't autofit)
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

    # tblGrid — authoritative column layout
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid_cols = "".join(f'<w:gridCol w:w="{w}"/>' for w in dxa)
    # Insert tblGrid right after tblPr
    grid = parse_xml(f'<w:tblGrid {nsdecls("w")}>{grid_cols}</w:tblGrid>')
    tblPr.addnext(grid)


def _set_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    ))


def _set_cell_style(cell, style_name, text, alignment,
                    fill="FFFFFF", top="none", bottom="none", left="none", right="none"):
    """Set cell paragraph style, text, alignment, shading and borders."""
    # Clear existing content and set style + text
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
    tc.append(p)

    # Apply paragraph style
    pPr = parse_xml(f'<w:pPr {nsdecls("w")}><w:pStyle w:val="{_style_id(style_name)}"/></w:pPr>')
    p.append(pPr)

    # Alignment
    jc_map = {
        WD_ALIGN_PARAGRAPH.LEFT:    "left",
        WD_ALIGN_PARAGRAPH.RIGHT:   "right",
        WD_ALIGN_PARAGRAPH.CENTER:  "center",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    }
    if alignment in jc_map:
        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="{jc_map[alignment]}"/>')
        pPr.append(jc)

    # Text run
    if text:
        r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{_xml_escape(text)}</w:t></w:r>')
        p.append(r)

    # Cell properties
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)

    # Shading
    existing_shd = tcPr.find(qn("w:shd"))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>'))

    # Borders
    existing_borders = tcPr.find(qn("w:tcBorders"))
    if existing_borders is not None:
        tcPr.remove(existing_borders)

    def bxml(name, val):
        if val == "none":
            return f'<w:{name} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        style, sz, color = val
        return f'<w:{name} w:val="{style}" w:sz="{sz}" w:space="0" w:color="{color}"/>'

    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'{bxml("top", top)}{bxml("bottom", bottom)}'
        f'{bxml("left", left)}{bxml("right", right)}'
        f'</w:tcBorders>'
    ))

    # Cell margins (2mm top/bottom, 1.5mm left/right — matches master template)
    existing_mar = tcPr.find(qn("w:tcMar"))
    if existing_mar is not None:
        tcPr.remove(existing_mar)
    tcPr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="113" w:type="dxa"/>'
        f'<w:bottom w:w="113" w:type="dxa"/>'
        f'<w:left w:w="85" w:type="dxa"/>'
        f'<w:right w:w="85" w:type="dxa"/>'
        f'</w:tcMar>'
    ))


def _style_id(style_name):
    """Convert style name to Word style ID (remove spaces)."""
    return style_name.replace(" ", "")


def _fmt_fee(fee_str):
    """
    Normalise simple fee strings: '4500' or '4,500' → '$4,500'.
    Complex strings (already formatted, with parentheses, multiple $ etc.) are
    returned as-is to avoid mangling compound values like '$2,000 (u.p. $3,500)'.
    """
    if not fee_str:
        return ""
    s = str(fee_str).strip()
    # If complex (parentheses, separators, multiple $, or narrative text after
    # the dollar amount) — return as-is to avoid mangling.
    # Check for "•" (U+2022 bullet) AND "·" (U+00B7 middle dot) — both appear
    # in Notion fee cells like "$3,000 • technicals at cost".
    if any(c in s for c in ("(", "/", "·", "•", "+", "setup")) or s.count("$") > 1:
        return s
    # String starts with $ and has a space → narrative text, leave untouched
    if s.startswith("$") and " " in s:
        return s
    # Strip spaces only for bare numeric values like "4 500" or "$ 4,500"
    s_clean = s.replace(" ", "")
    if s_clean.startswith("$"):
        return s_clean
    try:
        val = float(s_clean.replace(",", ""))
        return f"${val:,.0f}"
    except ValueError:
        return s


# ── Markdown parser ──────────────────────────────────────────────────────────

def generate_from_markdown(markdown: str, metadata: dict = None, filename: str = None) -> str:
    """
    PRIMARY API for skill invocations.

    Parse raw markdown and generate a branded .docx. Claude passes the
    markdown content as-is; this function handles all conversion so the
    output is always consistent.

    Parameters
    ----------
    markdown : str
        Raw markdown content (body only — no title/subtitle needed here).
    metadata : dict, optional
        Document metadata — all keys optional:
          doc_title, doc_subtitle, doc_type,
          recipient_salutation, recipient_name,
          signatory_name, signatory_title, filename
    filename : str, optional
        Override the output filename (without extension).

    Returns
    -------
    str
        Absolute path to the generated .docx file.

    Example
    -------
        path = generate_from_markdown(markdown_text, {
            "doc_title":   "Service Proposal",
            "doc_subtitle": "TRINITY CONSULTING",
            "recipient_salutation": "Mr",
            "recipient_name": "Clarence Loh",
            "signatory_name": "Michael Ryan Chan",
            "signatory_title": "Managing Partner",
        })
    """
    import re
    meta = metadata or {}
    blocks = _parse_markdown(markdown)

    # ── Remove leading salutation paragraph if recipient is also in metadata ──
    # Prevents "Dear Mr X," appearing twice (once from metadata, once from markdown body)
    if blocks and meta.get("recipient_name"):
        first = blocks[0]
        if first.get("type") == "paragraph" and re.match(r'^dear\b', first.get("text",""), re.I):
            blocks.pop(0)
            # Also remove the spacer that follows it
            if blocks and blocks[0].get("type") == "spacer":
                blocks.pop(0)

    # ── Smart spacer cleanup ──────────────────────────────────────────────────
    # Word styles have built-in space_after, so explicit spacer paragraphs
    # are only needed at section boundaries. Remove spacers elsewhere to avoid
    # double-height gaps between regular paragraphs.
    blocks = _clean_spacers(blocks)

    data = {
        "doc_title":             meta.get("doc_title", ""),
        "doc_subtitle":          meta.get("doc_subtitle", ""),
        "recipient_salutation":  meta.get("recipient_salutation", ""),
        "recipient_name":        meta.get("recipient_name", ""),
        "body":                  blocks,
        "signatory_name":        meta.get("signatory_name", "Michael Ryan Chan"),
        "signatory_title":       meta.get("signatory_title", "Managing Partner"),
    }
    doc_type = meta.get("doc_type", "letter")
    return generate_docx(doc_type, data, filename or meta.get("filename"))


def generate_from_notion_page(page: dict, metadata_overrides: dict = None, filename: str = None) -> str:
    """
    Generate a .docx directly from a notion-fetch result dict.

    Parameters
    ----------
    page : dict
        The dict returned by the notion-fetch MCP tool.
    metadata_overrides : dict, optional
        Override any auto-extracted metadata (doc_title, doc_subtitle,
        recipient_salutation, recipient_name, signatory_name, signatory_title).
    filename : str, optional
        Output filename (no extension).

    Returns
    -------
    str  — absolute path to the generated .docx
    """
    import re
    raw_text  = page.get("text", "")
    page_title = page.get("title", "")

    # Extract <content>…</content> block
    m = re.search(r'<content>(.*?)</content>', raw_text, re.S)
    content = m.group(1).strip() if m else raw_text

    # ── Auto-extract metadata from content ─────────────────────
    meta = {}

    # Page title (strip emoji prefix)
    clean_title = re.sub(r'^[\U00010000-\U0010ffff\U0001F300-\U0001FAFF\U00002600-\U000027BF\s]+', '', page_title).strip()
    # Split on " — " or " - " to separate parts.
    # Convention: "Doc Type — Client Name"  e.g. "Service Proposal — Trinity Consulting"
    # The LAST part is the client name (subtitle); the first part is the document type (title).
    if " — " in clean_title:
        parts = clean_title.split(" — ", 1)
        meta["doc_title"]    = parts[0].strip()
        meta["doc_subtitle"] = parts[1].strip().upper()
    elif " - " in clean_title:
        parts = clean_title.split(" - ", 1)
        meta["doc_title"]    = parts[0].strip()
        meta["doc_subtitle"] = parts[1].strip().upper()
    else:
        meta["doc_title"] = clean_title

    # Detect recipient from "Dear Mr/Ms/Dr X," opening line
    dear_m = re.search(r'Dear\s+(Mr|Ms|Mrs|Dr|Prof)\.?\s+([\w\s]+?)(?:,|\n)', content)
    if dear_m:
        meta["recipient_salutation"] = dear_m.group(1)
        meta["recipient_name"]       = dear_m.group(2).strip()

    # Detect signatory from "Warm regards," / "Yours sincerely," block
    sig_m = re.search(r'(?:Warm regards|Yours sincerely|Regards)[,\n]+\s*\*?\*?([\w\s]+?)\*?\*?\n([\w\s,]+?)(?:\n|$)', content, re.I)
    if sig_m:
        meta["signatory_name"]  = sig_m.group(1).strip()
        meta["signatory_title"] = sig_m.group(2).strip().rstrip(",")

    # Defaults
    meta.setdefault("signatory_name",  "Michael Ryan Chan")
    meta.setdefault("signatory_title", "Managing Partner")

    # Apply overrides
    if metadata_overrides:
        meta.update(metadata_overrides)

    # ── Strip closing signature block from content ──────────────
    # Remove "Warm regards, / Name / Title / URL" — it's added by the template
    content = re.sub(
        r'\n(?:Warm regards|Yours sincerely|Regards)[,\n].*$', '', content, flags=re.S | re.I
    ).strip()

    # ── Strip top metadata block (date line, bold subtitle, first ---) ─────
    # Notion pages often start with a date + client name before the first rule
    content = re.sub(r'^.*?---\n', '', content, count=1, flags=re.S).strip()

    # ── Preprocess Notion enhanced markdown → standard markdown ─
    content = _preprocess_notion_markdown(content)

    return generate_from_markdown(content, meta, filename)


def generate_from_notion_blocks(blocks: list, page_properties: dict = None,
                                metadata_overrides: dict = None, filename: str = None) -> str:
    """
    Generate a .docx directly from raw Notion API block objects.

    Unlike generate_from_notion_page (which relies on text already extracted
    by the notion-fetch MCP), this function works with the raw Notion REST API
    block format — used by the Evye bot which fetches blocks directly.

    Table blocks should have _rows pre-populated with their table_row children
    (Notion requires a separate /blocks/{tableId}/children call to get rows).

    Parameters
    ----------
    blocks : list
        Notion block objects from GET /v1/blocks/{id}/children.
        Each table block must have _rows: list of table_row objects.
    page_properties : dict, optional
        Full Notion page response from GET /v1/pages/{id}.
        Used to extract the page title for doc_title / doc_subtitle.
    metadata_overrides : dict, optional
        Override any auto-extracted metadata (doc_title, doc_subtitle,
        recipient_salutation, recipient_name, signatory_name, signatory_title).
    filename : str, optional
        Output filename (no extension).

    Returns
    -------
    str — absolute path to the generated .docx
    """
    import re

    # ── Extract metadata from page title ─────────────────────────
    meta = {}
    page_title = _extract_notion_page_title(page_properties or {})
    if page_title:
        clean_title = re.sub(
            r'^[\U00010000-\U0010ffff\U0001F300-\U0001FAFF\U00002600-\U000027BF\s]+',
            '', page_title
        ).strip()
        if " — " in clean_title:
            parts = clean_title.split(" — ", 1)
            meta["doc_title"]    = parts[0].strip()
            meta["doc_subtitle"] = parts[1].strip().upper()
        elif " - " in clean_title:
            parts = clean_title.split(" - ", 1)
            meta["doc_title"]    = parts[0].strip()
            meta["doc_subtitle"] = parts[1].strip().upper()
        else:
            meta["doc_title"] = clean_title

    meta.setdefault("signatory_name",  "Michael Ryan Chan")
    meta.setdefault("signatory_title", "Managing Partner")

    # ── Convert Notion blocks → content blocks + extract inline meta ──
    content_blocks, extracted_meta = _notion_blocks_to_content(blocks)
    meta.update(extracted_meta)

    # ── Title fallback: if page_properties had no title, promote the
    # first heading in the body (if it's near the top) to doc_title.
    if not meta.get("doc_title"):
        for i, b in enumerate(content_blocks[:3]):
            t = b.get("type")
            if t in ("heading2", "heading3"):
                meta["doc_title"] = b["text"]
                content_blocks.pop(i)
                break
            if t in ("paragraph", "bullet", "fee_table", "kv_table", "process_table"):
                # Non-heading content came first — give up, leave body intact.
                break

    # Apply caller overrides last (non-empty values only)
    if metadata_overrides:
        meta.update({k: v for k, v in metadata_overrides.items() if v})

    # ── Strip "Dear X" paragraph if recipient extracted (avoids duplication) ──
    if content_blocks and meta.get("recipient_name"):
        first = content_blocks[0]
        if first.get("type") == "paragraph" and re.match(r'^dear\b', first.get("text", ""), re.I):
            content_blocks.pop(0)
            if content_blocks and content_blocks[0].get("type") == "spacer":
                content_blocks.pop(0)

    # ── Spacer cleanup ────────────────────────────────────────────
    content_blocks = _clean_spacers(content_blocks)

    data = {
        "doc_title":            meta.get("doc_title", ""),
        "doc_subtitle":         meta.get("doc_subtitle", ""),
        "recipient_salutation": meta.get("recipient_salutation", ""),
        "recipient_name":       meta.get("recipient_name", ""),
        "body":                 content_blocks,
        "signatory_name":       meta.get("signatory_name", "Michael Ryan Chan"),
        "signatory_title":      meta.get("signatory_title", "Managing Partner"),
    }
    return generate_docx("letter", data, filename or meta.get("filename"))


def _extract_notion_page_title(page_obj: dict) -> str:
    """
    Extract plain-text title from a Notion page object (full GET /pages/{id} response)
    or a bare properties dict.
    """
    if not page_obj:
        return ""
    # Full page object has a nested "properties" key
    properties = page_obj.get("properties") if "properties" in page_obj else page_obj
    if not isinstance(properties, dict):
        return ""

    # Notion: exactly one property per page/DB entry has type == "title".
    # Its key varies ("Name", "Task name", "Title", etc.) so scan by type.
    for prop in properties.values():
        if isinstance(prop, dict) and prop.get("type") == "title":
            text = _extract_notion_rich_text(prop.get("title") or [])
            if text:
                return text

    # Fallback: common key names (bare properties dict without type info)
    for key in ("title", "Title", "Name", "name", "Task name", "Task Name"):
        prop = properties.get(key)
        if not prop:
            continue
        if isinstance(prop, dict):
            rich_text = prop.get("title") or prop.get("rich_text") or []
            text = _extract_notion_rich_text(rich_text)
            if text:
                return text
        elif isinstance(prop, list):
            text = _extract_notion_rich_text(prop)
            if text:
                return text
    return ""


def _extract_notion_rich_text(rich_text: list) -> str:
    """
    Concatenate plain text from a Notion rich_text array.
    Uses the convenience plain_text field Notion always provides.
    """
    if not rich_text:
        return ""
    parts = []
    for rt in rich_text:
        text = rt.get("plain_text") or rt.get("text", {}).get("content", "")
        parts.append(text)
    return "".join(parts).strip()


def _extract_notion_runs(rich_text: list) -> list:
    """
    Convert a Notion rich_text array into renderer run dicts, preserving
    bold/italic/underline/strikethrough/code annotations and links.
    """
    runs = []
    for rt in (rich_text or []):
        text = rt.get("plain_text") or rt.get("text", {}).get("content", "")
        if not text:
            continue
        ann = rt.get("annotations") or {}
        runs.append({
            "t": text,
            "b": bool(ann.get("bold")),
            "i": bool(ann.get("italic")),
            "u": bool(ann.get("underline")),
            "s": bool(ann.get("strikethrough")),
            "c": bool(ann.get("code")),
            "link": rt.get("href") or None,
        })
    return runs


def _get_notion_block_text(block: dict) -> str:
    """Return plain text for any Notion block (used for preamble/sign-off detection)."""
    btype = block.get("type", "")
    return _extract_notion_rich_text(block.get(btype, {}).get("rich_text", []))


def _notion_blocks_to_content(blocks: list) -> tuple:
    """
    Convert Notion API block objects to (content_blocks, meta).

    Philosophy: trust the content. Pass everything through by default.
    The only rewrite is at the very end of the document: if the author
    typed their own sign-off block (Warm regards / Yours sincerely / etc.)
    followed by a name line, promote it to signatory metadata and remove
    ONLY that trailing block — so the template renders it consistently
    without duplication. Nothing before the sign-off is ever touched.

    No preamble stripping. No content-based filtering. Dividers pass
    through as spacers and are rendered as-is.
    """
    import re

    meta = {}

    # ── Look for a sign-off block near the end of the document ─────────
    # Match only canonical letter sign-offs on their own line (full match).
    # Only inspect the LAST 8 non-blank blocks to avoid false positives from
    # "regards to X" style phrases appearing mid-body.
    SIGN_OFF = re.compile(
        r'^(talk soon|warm regards|yours sincerely|yours faithfully|'
        r'regards|best regards|kind regards|sincerely)[,\.!]?\s*$',
        re.I,
    )

    # Build an index of non-blank blocks
    non_blank_indices = [
        i for i, b in enumerate(blocks)
        if _get_notion_block_text(b).strip()
    ]

    sign_off_idx = None
    # Only scan the last 8 non-blank blocks — sign-off is always near the end
    for i in non_blank_indices[-8:]:
        text = _get_notion_block_text(blocks[i]).strip()
        if SIGN_OFF.match(text):
            sign_off_idx = i
            break

    # Require at least one non-blank block AFTER the sign-off that looks
    # like a name (short, title-cased, no sentence punctuation). Otherwise,
    # it's probably body text and we should leave it alone.
    if sign_off_idx is not None:
        tail_texts = [
            _get_notion_block_text(b).strip()
            for b in blocks[sign_off_idx + 1:]
            if _get_notion_block_text(b).strip()
        ]
        looks_like_name = bool(
            tail_texts
            and len(tail_texts[0]) <= 60
            and not tail_texts[0].endswith(('.', '?', '!'))
        )
        if looks_like_name:
            # First line = name; remaining non-URL lines = title/company
            filtered = [t for t in tail_texts if not re.match(r'https?://', t)]
            if filtered:
                meta["signatory_name"] = filtered[0]
            if len(filtered) > 1:
                title = filtered[1]
                # Strip trailing ", Evye LLP" / ", Evye Pte Ltd" so the
                # template's hardcoded "Evye LLP" line doesn't duplicate
                title = re.sub(
                    r',\s*Evye\s*(LLP|Pte\.?\s*Ltd|Ltd)?\.?\s*$',
                    '', title, flags=re.I
                ).strip()
                meta["signatory_title"] = title
            blocks = blocks[:sign_off_idx]
        # else: false positive — leave content alone

    # ── 3. Convert blocks ─────────────────────────────────────────
    # _emit_block handles a single Notion block + its children recursively.
    # Children appear on `blk["children"]` when the n8n workflow has
    # pre-fetched them (Notion API requires a separate GET per parent block
    # because GET /blocks/{id}/children only returns one level deep).
    #
    # Bullets nest by incrementing `level` (capped at 3 to match Word styles).
    # Other block types (paragraph, heading, toggle, etc.) recurse with the
    # same level — children render inline beneath their parent.
    def _emit_block(blk, level=1):
        btype = blk.get("type", "")
        children = blk.get("children", []) or []
        out = []

        def _payload():
            return blk.get(btype, {}) or {}

        if btype == "paragraph":
            runs = _extract_notion_runs(_payload().get("rich_text", []))
            text = _runs_plain(runs)
            if not text:
                out.append({"type": "spacer"})
            else:
                if level == 1 and not meta.get("recipient_name"):
                    dear_m = re.match(
                        r'^[Dd]ear\s+(?:(Mr|Ms|Mrs|Dr|Prof)\.?\s+)?([\w][\w\s]*?)(?:,\s*)?$',
                        text
                    )
                    if dear_m:
                        if dear_m.group(1):
                            meta["recipient_salutation"] = dear_m.group(1)
                            meta["recipient_name"] = dear_m.group(2).strip()
                        else:
                            meta["recipient_name"] = dear_m.group(2).strip()
                out.append({"type": "paragraph", "text": text, "runs": runs})

        elif btype == "heading_1":
            text = _extract_notion_rich_text(_payload().get("rich_text", []))
            if text:
                out.append({"type": "heading2", "text": text})

        elif btype == "heading_2":
            text = _extract_notion_rich_text(_payload().get("rich_text", []))
            if text:
                out.append({"type": "heading3", "text": text})

        elif btype == "heading_3":
            text = _extract_notion_rich_text(_payload().get("rich_text", []))
            if text:
                out.append({"type": "heading4", "text": text})

        elif btype == "bulleted_list_item":
            runs = _extract_notion_runs(_payload().get("rich_text", []))
            text = _runs_plain(runs)
            if text:
                out.append({"type": "bullet", "text": text, "runs": runs,
                            "level": min(level, 3)})
            for child in children:
                out.extend(_emit_block(child, level + 1))
            return out

        elif btype == "numbered_list_item":
            runs = _extract_notion_runs(_payload().get("rich_text", []))
            text = _runs_plain(runs)
            if text:
                out.append({"type": "numbered", "text": text, "runs": runs,
                            "level": min(level, 3)})
            for child in children:
                out.extend(_emit_block(child, level + 1))
            return out

        elif btype == "to_do":
            runs = _extract_notion_runs(_payload().get("rich_text", []))
            text = _runs_plain(runs)
            if text:
                out.append({"type": "todo", "checked": bool(_payload().get("checked")),
                            "text": text, "runs": runs, "level": min(level, 3)})
            for child in children:
                out.extend(_emit_block(child, level + 1))
            return out

        elif btype == "toggle":
            runs = _extract_notion_runs(_payload().get("rich_text", []))
            text = _runs_plain(runs)
            if text:
                out.append({"type": "paragraph", "text": text, "runs": runs})
            for child in children:
                out.extend(_emit_block(child, level))
            return out

        elif btype in ("quote", "callout"):
            runs = _extract_notion_runs(_payload().get("rich_text", []))
            text = _runs_plain(runs)
            if text:
                out.append({"type": "paragraph_left", "text": text, "runs": runs})

        elif btype == "divider":
            out.append({"type": "spacer"})

        elif btype == "table":
            table_block = _convert_notion_table_block(blk)
            if table_block:
                out.append(table_block)
            return out

        elif btype == "code":
            code_text = _extract_notion_rich_text(_payload().get("rich_text", []))
            if code_text:
                out.append({"type": "code_block", "text": code_text})
            cap = _extract_notion_rich_text(_payload().get("caption", []))
            if cap:
                out.append({"type": "paragraph_left", "text": cap})

        elif btype == "image":
            img = _payload()
            url = ((img.get("file") or {}).get("url")
                   or (img.get("external") or {}).get("url"))
            caption = _extract_notion_rich_text(img.get("caption", []))
            if url:
                out.append({"type": "image", "url": url, "caption": caption})
            elif caption:
                out.append({"type": "paragraph_left", "text": caption})

        elif btype in ("bookmark", "embed", "link_preview"):
            data = _payload()
            url = data.get("url", "")
            cap = _extract_notion_rich_text(data.get("caption", []))
            label = cap or url
            if label:
                out.append({"type": "paragraph", "text": label,
                            "runs": [{"t": label, "link": url or None}]})

        elif btype in ("file", "pdf", "video", "audio"):
            data = _payload()
            name = data.get("name") or ""
            cap = _extract_notion_rich_text(data.get("caption", []))
            url = ((data.get("file") or {}).get("url")
                   or (data.get("external") or {}).get("url") or "")
            label = name or cap or url or btype
            out.append({"type": "paragraph_left", "text": f"[Attachment: {label}]"})

        elif btype == "equation":
            expr = _payload().get("expression", "")
            if expr:
                out.append({"type": "code_block", "text": expr})

        elif btype in ("child_page", "child_database"):
            title = _payload().get("title", "")
            if title:
                out.append({"type": "paragraph_left", "text": f"[Sub-page: {title}]"})

        elif btype in ("table_of_contents", "breadcrumb", "template"):
            pass  # purely navigational — nothing to render

        elif btype in ("column_list", "column", "synced_block"):
            pass  # containers — content is in children, recursed below

        else:
            # ── No-silent-loss fallback ──
            text = _get_notion_block_text(blk)
            if text:
                out.append({"type": "paragraph", "text": text})
            elif not children:
                out.append({"type": "paragraph_left",
                            "text": f"[Unsupported block: {btype}]"})

        # Default: recurse into children for any type that didn't return above
        for child in children:
            out.extend(_emit_block(child, level))
        return out

    content_blocks = []
    for blk in blocks:
        content_blocks.extend(_emit_block(blk))

    # ── Infer bullets from prose paragraphs ───────────────────────────
    # Writers often prefix paragraphs with "✦ ", "• ", "- " etc. instead
    # of using Notion's bulleted_list_item block. Detect runs of 2+
    # consecutive paragraphs starting with a bullet glyph and promote
    # them to bullet blocks. Keeps single-glyph decorative lines alone.
    BULLET_PREFIX = re.compile(r'^\s*([\u2726\u2727\u2022\u00b7\u25aa\u25a0\u2043\u2013\u2014*\-])\s+(.+)$')
    i = 0
    while i < len(content_blocks):
        blk = content_blocks[i]
        if blk.get("type") == "paragraph":
            m = BULLET_PREFIX.match(blk.get("text", ""))
            if m:
                # Look ahead for a run
                run_end = i
                while run_end < len(content_blocks):
                    nb = content_blocks[run_end]
                    if nb.get("type") != "paragraph":
                        break
                    nm = BULLET_PREFIX.match(nb.get("text", ""))
                    if not nm:
                        break
                    run_end += 1
                run_len = run_end - i
                if run_len >= 2:
                    for j in range(i, run_end):
                        blk_j = content_blocks[j]
                        m2 = BULLET_PREFIX.match(blk_j["text"])
                        stripped_text = m2.group(2).strip()
                        new_blk = {"type": "bullet", "text": stripped_text, "level": 1}
                        runs = blk_j.get("runs")
                        if runs:
                            remaining = m2.start(2)   # chars to strip from run stream
                            new_runs = []
                            for r in runs:
                                t = r.get("t", "")
                                if remaining > 0:
                                    if len(t) <= remaining:
                                        remaining -= len(t)
                                        continue
                                    t = t[remaining:]
                                    remaining = 0
                                new_runs.append({**r, "t": t})
                            new_blk["runs"] = new_runs
                        content_blocks[j] = new_blk
                    i = run_end
                    continue
        i += 1

    return content_blocks, meta


def _convert_notion_table_block(block: dict) -> dict:
    """
    Convert a Notion table block (with _rows pre-fetched) to a
    fee_table, kv_table, or process_table content block.
    Auto-detects table type from column count and header text.
    """
    import re

    table_meta = block.get("table", {})
    has_column_header = table_meta.get("has_column_header", False)
    rows = block.get("_rows", [])
    if not rows:
        return None

    # Parse each row's cells
    parsed_rows = []
    for row_blk in rows:
        if row_blk.get("type") != "table_row":
            continue
        cells = row_blk.get("table_row", {}).get("cells", [])
        parsed_rows.append([_extract_notion_rich_text(cell) for cell in cells])

    if not parsed_rows:
        return None

    n_cols = max(len(r) for r in parsed_rows)

    # Separate header row
    headers = None
    data_rows = parsed_rows
    if has_column_header and len(parsed_rows) > 1:
        headers = parsed_rows[0]
        data_rows = parsed_rows[1:]

    if n_cols == 2:
        rows_out = [
            {"label": r[0] if len(r) > 0 else "", "value": r[1] if len(r) > 1 else ""}
            for r in data_rows
        ]
        return {"type": "kv_table", "rows": rows_out}

    if n_cols == 3:
        FEE_WORDS = {'fee', 'investment', 'cost', 'price', 'amount', 'rate',
                     'total', 'charge', 'quoted', 'sgd', 's$', '$'}
        last_header = (headers[-1] if headers else "").lower()
        is_fee = any(w in last_header for w in FEE_WORDS)
        if not is_fee and data_rows:
            last_vals = [r[-1] if r else "" for r in data_rows[:3]]
            is_fee = any(re.search(r'[\$\d,]{3,}', v) for v in last_vals)

        if is_fee:
            rows_out = []
            total = None
            for r in data_rows:
                item = r[0] if len(r) > 0 else ""
                desc = r[1] if len(r) > 1 else ""
                fee  = r[2] if len(r) > 2 else ""
                if re.match(r'^total', item.lower()) or (not item and re.match(r'^total', desc.lower())):
                    total = fee
                else:
                    rows_out.append({"item": item, "description": desc, "fee": fee})
            return {
                "type": "fee_table",
                "headers": headers or ["Item", "Description", "Fee"],
                "rows": rows_out,
                "total": total,
            }

        rows_out = [
            {"col1": r[0] if len(r) > 0 else "",
             "col2": r[1] if len(r) > 1 else "",
             "col3": r[2] if len(r) > 2 else ""}
            for r in data_rows
        ]
        return {"type": "process_table", "headers": headers, "rows": rows_out}

    # 1 column or 4+ columns → generic table. NOTHING is dropped or truncated.
    return {"type": "generic_table", "headers": headers, "rows": data_rows}


def _preprocess_notion_markdown(text: str) -> str:
    """
    Convert Notion enhanced markdown (as returned by notion-fetch MCP) to
    standard markdown that _parse_markdown can handle.

    Handles:
      - <table> HTML blocks → pipe markdown tables
      - <callout icon="…"> blocks → > blockquote
      - <span color="…">text</span> → text
      - Escaped dollar signs \\$ → $
      - <br> → separator (· between fee amounts)
      - HTML comments / leftover tags → stripped
    """
    import re

    # 1. Unescape dollar signs
    text = text.replace('\\$', '$')

    # 2. Strip <span color="...">text</span> — keep inner text
    text = re.sub(r'<span[^>]*>(.*?)</span>', r'\1', text, flags=re.S)

    # 3. Convert <callout …>…</callout> → structured markdown
    #    Preserves: bold title → heading4; body paragraphs; bullet lists.
    #    Notion indents callout inner content with tabs — strip them.
    def convert_callout(m):
        inner = m.group(1)
        # Strip one level of leading tab/spaces from each line
        lines = [re.sub(r'^[\t ]{1,4}', '', ln) for ln in inner.splitlines()]
        # Drop leading/trailing blank lines
        while lines and not lines[0].strip():
            lines.pop(0)
        while lines and not lines[-1].strip():
            lines.pop()

        out = []
        # First line, if fully bold, becomes a heading4 (bold italic)
        if lines and re.match(r'^\*\*(.+)\*\*\s*$', lines[0].strip()):
            title = re.match(r'^\*\*(.+)\*\*\s*$', lines[0].strip()).group(1)
            out.append(f'#### {title}')
            lines.pop(0)
        # Drop blank line right after the title (the parser will add spacing)
        while lines and not lines[0].strip():
            lines.pop(0)

        out.append('')  # blank line before body
        out.extend(lines)
        out.append('')  # blank line after
        return '\n'.join(out)
    text = re.sub(r'<callout[^>]*>(.*?)</callout>', convert_callout, text, flags=re.S)

    # 4. Convert <table>…</table> → pipe markdown table
    def convert_table(m):
        table_html = m.group(0)
        rows = re.findall(r'<tr>(.*?)</tr>', table_html, re.S)
        md_rows = []
        for ri, row in enumerate(rows):
            cells = re.findall(r'<td[^>]*>(.*?)</td>', row, re.S)
            # Clean each cell: strip tags, bold markers, join <br> with ·
            clean_cells = []
            for cell in cells:
                cell = re.sub(r'<br\s*/?>', ' · ', cell)
                cell = re.sub(r'<[^>]+>', '', cell)           # strip remaining HTML
                cell = re.sub(r'\*\*(.+?)\*\*', r'\1', cell) # strip bold
                cell = re.sub(r'\*(.+?)\*', r'\1', cell)      # strip italic
                cell = cell.strip()
                clean_cells.append(cell)
            md_rows.append('| ' + ' | '.join(clean_cells) + ' |')
            if ri == 0:
                md_rows.append('| ' + ' | '.join(['---'] * len(clean_cells)) + ' |')
        return '\n'.join(md_rows)
    text = re.sub(r'<table[^>]*>.*?</table>', convert_table, text, flags=re.S)

    # 5. Strip any remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)

    # 6. Collapse 3+ blank lines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 7. Notion uses single newlines as paragraph breaks, but markdown needs
    #    blank lines. Insert blank lines between consecutive prose lines.
    def is_prose(line):
        s = line.strip()
        if not s:
            return False
        # Non-prose markers: headings, bullets, rules, blockquotes, tables
        return not re.match(r'^(#{1,6}\s|[-*+]\s|✦\s|•\s|\d+[.)] |>|\||---|\s*\|)', s)

    lines = text.split('\n')
    out = []
    for i, line in enumerate(lines):
        out.append(line)
        if i + 1 < len(lines) and is_prose(line) and is_prose(lines[i + 1]):
            out.append('')
    text = '\n'.join(out)

    return text.strip()


def _clean_spacers(blocks: list) -> list:
    """
    Keep spacers only at section boundaries. Remove them everywhere else
    so Word's built-in paragraph spacing handles normal content flow.

    KEEP spacer when:
      - before a heading (section break)
      - after a table (visual separation)
      - before a table (visual separation)
      - before/after a paragraph_left (aside/blockquote)

    REMOVE spacer when:
      - between paragraph → paragraph
      - between paragraph → bullet
      - between bullet → bullet
      - between bullet → paragraph
    """
    SECTION_TYPES = {"heading2", "heading3", "heading4"}
    TABLE_TYPES   = {"fee_table", "kv_table", "process_table", "generic_table"}
    ASIDE_TYPES   = {"paragraph_left"}

    def keep_spacer(prev_type, next_type):
        if prev_type is None or next_type is None:
            return False
        # Always keep before/after a heading
        if next_type in SECTION_TYPES:
            return True
        # Always keep after a table
        if prev_type in TABLE_TYPES:
            return True
        # Always keep before a table
        if next_type in TABLE_TYPES:
            return True
        # Keep around asides
        if prev_type in ASIDE_TYPES or next_type in ASIDE_TYPES:
            return True
        # Remove everywhere else (paragraph↔paragraph, paragraph↔bullet, bullet↔bullet)
        return False

    cleaned = []
    for j, blk in enumerate(blocks):
        if blk.get("type") == "spacer":
            prev_type = cleaned[-1].get("type") if cleaned else None
            next_type = blocks[j + 1].get("type") if j + 1 < len(blocks) else None
            if not keep_spacer(prev_type, next_type):
                continue
        cleaned.append(blk)

    # Collapse consecutive spacers to one
    deduped = []
    for blk in cleaned:
        if blk.get("type") == "spacer" and deduped and deduped[-1].get("type") == "spacer":
            continue
        deduped.append(blk)

    return deduped


def _parse_markdown(text: str) -> list:
    """
    Convert markdown text into a list of content blocks understood by the
    generator. Handles:
      - Headings (#/##/###/####) — strips leading section numbers
      - Bullet lists (-, *, +, ✦, •) with up to 3 indent levels
      - Markdown tables (auto-detects fee vs kv table)
      - Horizontal rules (---) → spacer
      - Aside/callout blocks (> text) → Body Left
      - Blank lines between sections → spacer (deduped)
      - Inline bold/italic markers are stripped (text is preserved)
    """
    import re

    lines = text.split("\n")
    blocks = []
    i = 0
    prev_was_spacer = False

    def add_spacer():
        nonlocal prev_was_spacer
        if not prev_was_spacer and blocks:
            blocks.append({"type": "spacer"})
            prev_was_spacer = True

    def add_block(b):
        nonlocal prev_was_spacer
        blocks.append(b)
        prev_was_spacer = False

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # ── Blank line ────────────────────────────────────────
        if not line:
            add_spacer()
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────
        if re.match(r'^[-_*]{3,}$', line):
            add_spacer()
            i += 1
            continue

        # ── Fenced code block ─────────────────────────────────
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if code_lines:
                add_block({"type": "code_block", "text": "\n".join(code_lines)})
            continue

        # ── Heading ───────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading_text = _strip_inline(m.group(2))
            # Strip leading section numbers: "1. Title" or "1) Title"
            heading_text = re.sub(r'^\d+[.)]\s+', '', heading_text).strip()
            type_map = {1: "heading2", 2: "heading2", 3: "heading3", 4: "heading4"}
            add_block({"type": type_map.get(level, "heading2"), "text": heading_text})
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            block = _parse_table(table_lines)
            if block:
                add_block(block)
            continue

        # ── Blockquote / aside ────────────────────────────────
        if line.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            text = _strip_inline(' '.join(quote_lines))
            add_block({"type": "paragraph_left", "text": text,
                       "runs": _parse_inline(' '.join(quote_lines))})
            continue

        # ── Todo item: "- [ ] text" / "- [x] text" ────────────
        m = re.match(r'^(\s*)[-*+] \[( |x|X)\] (.+)', raw)
        if m:
            indent = len(m.group(1))
            level = 1 if indent < 2 else (2 if indent < 4 else 3)
            add_block({"type": "todo", "checked": m.group(2).lower() == "x",
                       "text": _strip_inline(m.group(3)),
                       "runs": _parse_inline(m.group(3)), "level": level})
            i += 1
            continue

        # ── Numbered list item: "1. text" / "2) text" ─────────
        m = re.match(r'^(\s*)\d+[.)] (.+)', raw)
        if m:
            indent = len(m.group(1))
            level = 1 if indent < 2 else (2 if indent < 4 else 3)
            add_block({"type": "numbered", "text": _strip_inline(m.group(2)),
                       "runs": _parse_inline(m.group(2)), "level": level})
            i += 1
            continue

        # ── Bullet point ──────────────────────────────────────
        # Matches: "  - text", "  * text", "  + text", "✦ text", "• text"
        m = re.match(r'^(\s*)([-*+]|✦|•) (.+)', raw)
        if m:
            indent = len(m.group(1))
            text = _strip_inline(m.group(3))
            level = 1 if indent < 2 else (2 if indent < 4 else 3)
            add_block({"type": "bullet", "text": text,
                       "runs": _parse_inline(m.group(3)), "level": level})
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────
        # Collect lines until a blank line, heading, table, or rule
        para_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                break
            if re.match(r'^#{1,4}\s', l):
                break
            if l.startswith('|'):
                break
            if re.match(r'^[-_*]{3,}$', l):
                break
            if re.match(r'^(\s*)([-*+]|✦|•|\d+[.)]) ', lines[i]):
                break
            if l.startswith('>'):
                break
            if l.startswith('```'):
                break
            para_lines.append(l)
            i += 1

        if para_lines:
            joined = ' '.join(para_lines)
            text = _strip_inline(joined)
            if text:
                add_block({"type": "paragraph", "text": text,
                           "runs": _parse_inline(joined)})
        continue

    # Trim leading/trailing spacers
    while blocks and blocks[0].get("type") == "spacer":
        blocks.pop(0)
    while blocks and blocks[-1].get("type") == "spacer":
        blocks.pop()

    return blocks


def _parse_table(lines: list) -> dict:
    """Parse markdown table lines into a fee_table or kv_table block."""
    import re

    # Filter separator rows (|---|---|)
    data_rows = [l for l in lines if not re.match(r'^\|[-|\s:]+\|$', l)]
    if not data_rows:
        return None

    def parse_cells(row):
        return [_strip_inline(c.strip()) for c in row.strip('|').split('|')]

    parsed = [parse_cells(row) for row in data_rows]
    if not parsed:
        return None

    headers = parsed[0]
    data    = parsed[1:]
    n_cols  = len(headers)

    if n_cols == 2:
        # Key-value table
        rows = [{"label": r[0] if len(r) > 0 else "",
                 "value": r[1] if len(r) > 1 else ""}
                for r in data]
        return {"type": "kv_table", "rows": rows}

    elif n_cols >= 3:
        # Determine if this is a fee table or a process/step table
        FEE_WORDS = {'fee', 'investment', 'cost', 'price', 'amount', 'rate', 'total', 'charge', 'quoted'}
        third_header = headers[2].lower().strip()
        is_fee = any(word in third_header for word in FEE_WORDS)

        if is_fee:
            # Fee table — Item | Description | Fee
            rows  = []
            total = None
            for row in data:
                item = row[0] if len(row) > 0 else ""
                desc = row[1] if len(row) > 1 else ""
                fee  = row[2] if len(row) > 2 else ""
                if re.match(r'^total', item.lower()) or (not item and re.match(r'^total', desc.lower())):
                    total = fee
                else:
                    rows.append({"item": item, "description": desc, "fee": fee})
            return {"type": "fee_table", "headers": headers, "rows": rows, "total": total}
        else:
            # Process/step table — 3 left-aligned columns with equal-ish widths
            rows = [{"col1": row[0] if len(row) > 0 else "",
                     "col2": row[1] if len(row) > 1 else "",
                     "col3": row[2] if len(row) > 2 else ""}
                    for row in data]
            return {"type": "process_table", "headers": headers, "rows": rows}

    return None


_INLINE_RE = None

def _parse_inline(text: str) -> list:
    """Tokenize markdown inline formatting into renderer run dicts."""
    import re
    global _INLINE_RE
    if _INLINE_RE is None:
        _INLINE_RE = re.compile(
            r'\*\*\*(?P<bi>.+?)\*\*\*'
            r'|\*\*(?P<b>.+?)\*\*'
            r'|__(?P<b2>.+?)__'
            r'|\*(?P<i>[^*\n]+?)\*'
            r'|(?<![\w])_(?P<i2>[^_\n]+?)_(?![\w])'
            r'|`(?P<c>[^`\n]+?)`'
            r'|\[(?P<lt>[^\]]+)\]\((?P<lu>[^)\s]+)\)'
        )
    runs, pos = [], 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            runs.append({"t": text[pos:m.start()]})
        if m.group("bi") is not None:
            runs.append({"t": m.group("bi"), "b": True, "i": True})
        elif m.group("b") is not None:
            runs.append({"t": m.group("b"), "b": True})
        elif m.group("b2") is not None:
            runs.append({"t": m.group("b2"), "b": True})
        elif m.group("i") is not None:
            runs.append({"t": m.group("i"), "i": True})
        elif m.group("i2") is not None:
            runs.append({"t": m.group("i2"), "i": True})
        elif m.group("c") is not None:
            runs.append({"t": m.group("c"), "c": True})
        elif m.group("lt") is not None:
            runs.append({"t": m.group("lt"), "link": m.group("lu")})
        pos = m.end()
    if pos < len(text):
        runs.append({"t": text[pos:]})
    return runs


def _strip_inline(text: str) -> str:
    """Remove markdown inline formatting markers, preserving the text."""
    import re
    # Bold-italic: ***text***
    text = re.sub(r'\*{3}(.+?)\*{3}', r'\1', text)
    # Bold: **text** or __text__
    text = re.sub(r'\*{2}(.+?)\*{2}', r'\1', text)
    text = re.sub(r'_{2}(.+?)_{2}',   r'\1', text)
    # Italic: *text* or _text_
    text = re.sub(r'\*(.+?)\*',  r'\1', text)
    text = re.sub(r'_(.+?)_',    r'\1', text)
    # Inline code
    text = re.sub(r'`(.+?)`',    r'\1', text)
    # Emoji/unicode symbols that are formatting-only (✦ as bullet marker handled elsewhere)
    text = re.sub(r'^✦\s*', '', text)
    text = re.sub(r'^•\s*',  '', text)
    # Trim
    return text.strip()


def _slugify(text):
    import re
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")[:40]


def _xml_escape(text):
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


# ── Quick smoke test ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    import docx as _docx_module
    import docx.enum.text

    path = generate_docx("letter", {
        "doc_title": "Letter of Engagement",
        "doc_subtitle": "BRAND IDENTITY PROJECT",
        "recipient_salutation": "Ms",
        "recipient_name": "Sarah Tan",
        "body": [
            {"type": "paragraph", "text": "We are pleased to confirm your engagement with Evye LLP for the Brand Identity project commencing 1 May 2026."},
            {"type": "heading2", "text": "Scope of Work"},
            {"type": "paragraph", "text": "The scope covers the following deliverables:"},
            {"type": "bullet", "text": "Brand strategy and positioning workshop", "level": 1},
            {"type": "bullet", "text": "Visual identity system (logo, palette, typography)", "level": 1},
            {"type": "bullet", "text": "Brand guidelines document", "level": 1},
            {"type": "heading2", "text": "Fee Summary"},
            {"type": "fee_table", "rows": [
                {"item": "Brand Strategy", "description": "Discovery, research, positioning workshop", "fee": "$4,500"},
                {"item": "Visual Identity", "description": "Logo design, colour palette, typography system", "fee": "$8,000"},
            ], "total": "$12,500"},
        ],
        "signatory_name": "Michael Ryan Chan",
        "signatory_title": "Managing Partner",
    })

    print(f"\nGenerated: {path}")
    import subprocess
    subprocess.run(["open", path])
