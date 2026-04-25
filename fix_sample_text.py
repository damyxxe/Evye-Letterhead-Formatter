"""
Update all sample text in evye-master.docx to accurately reflect
the current template styles and serve as a useful style guide.
"""

from pathlib import Path
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent
TEMPLATE_PATH = PROJECT_ROOT / "templates" / "evye-master.docx"


# Map of paragraph index → new text content
# (None = delete paragraph, keep style)
UPDATES = {
    # Para 3: Body text intro — fix font size reference
    3: (
        "We are pleased to present this document as a demonstration of the Evye master "
        "template styles. This paragraph uses the Evye Body style with justified alignment, "
        "Helvetica Neue at 10pt, and 1.35x line spacing. All body text in Evye documents "
        "should use this style."
    ),

    # Para 5: H2 description
    5: "Heading 2 uses Telegraf Bold at 14pt for section titles.",

    # Para 7: H3 description
    7: "Heading 3 uses Helvetica Neue Bold at 10pt in uppercase — visually matches clause headings.",

    # Para 9: H4 description
    9: "Heading 4 uses Helvetica Neue Bold Italic at 10pt for the finest heading level.",
}


def update_text(doc):
    print("Updating sample text...")
    for idx, new_text in UPDATES.items():
        p = doc.paragraphs[idx]
        old_text = p.text[:60]

        # Clear all runs and add new text
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)

        print(f"  Para {idx} [{p.style.name}]:")
        print(f"    was: {old_text}...")
        print(f"    now: {new_text[:60]}...")


def main():
    print(f"Opening: {TEMPLATE_PATH}")
    doc = Document(str(TEMPLATE_PATH))
    update_text(doc)
    doc.save(str(TEMPLATE_PATH))
    print(f"\nSaved: {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
