import sys, uuid
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO))


def doc_paragraphs(path):
    from docx import Document
    return [(p.style.name, p.text) for p in Document(str(path)).paragraphs]


def para_runs(path, index):
    from docx import Document
    p = Document(str(path)).paragraphs[index]
    out = []
    for r in p.runs:
        out.append({
            "text": r.text,
            "bold": bool(r.bold),
            "italic": bool(r.italic),
            "underline": bool(r.underline),
            "strike": bool(r.font.strike),
            "font": r.font.name,
        })
    return out


def doc_tables(path):
    from docx import Document
    doc = Document(str(path))
    return [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]


def image_count(path):
    from docx import Document
    return len(Document(str(path)).inline_shapes)


def hyperlinks(path):
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(path))
    out = []
    for p in doc.paragraphs:
        for h in p._p.findall(qn("w:hyperlink")):
            r_id = h.get(qn("r:id"))
            url = doc.part.rels[r_id].target_ref if r_id in doc.part.rels else ""
            text = "".join(t.text or "" for t in h.iter(qn("w:t")))
            out.append((text, url))
    return out


import pytest

@pytest.fixture
def gen():
    """Call word_generator entrypoints with a unique filename; delete after."""
    created = []

    def _run(fn, *args, **kwargs):
        kwargs.setdefault("filename", f"_pytest_{uuid.uuid4().hex[:8]}")
        path = fn(*args, **kwargs)
        created.append(path)
        return path

    yield _run
    for p in created:
        Path(p).unlink(missing_ok=True)
